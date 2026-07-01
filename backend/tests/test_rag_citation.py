import base64
import io

import pytest
from flask import Flask

from src.services.rag_citation_service import rag_citation_service


def test_verify_marks_unsupported_claims_and_degrades_when_required():
    content = "变量用于保存程序运行中的数据 [S1]。循环语句可以重复执行代码。"
    citations = [{"source_id": "S1", "title": "变量", "excerpt": "变量用于保存数据"}]

    report = rag_citation_service.verify(
        content,
        citations,
        rag_required=True,
        citation_style="inline",
    )

    assert report["status"] == "needs_review"
    assert report["degradation"] == "unsupported_claims"
    assert report["degraded"] is True
    assert report["rag_required"] is True
    assert report["citation_style"] == "inline"
    assert report["unsupported_claims"]
    assert any(issue["type"] == "unsupported_claims" for issue in report["citation_issues"])


def test_verify_accepts_non_s_source_ids():
    content = "变量用于保存程序运行中的数据 [KP12]。"
    citations = [{"source_id": "KP12", "title": "变量", "excerpt": "变量用于保存数据"}]

    report = rag_citation_service.verify(content, citations, rag_required=True)

    assert report["status"] == "passed"
    assert report["degradation"] is None
    assert report["citation_coverage_score"] == 100


def test_attach_citations_adds_reference_note_and_metadata():
    content = {"title": "变量说明", "body": "变量用于保存程序运行中的数据。"}
    evidence = [{"source_id": "S1", "source_type": "syllabus", "title": "变量", "excerpt": "变量用于保存数据"}]

    result = rag_citation_service.attach_citations(
        content,
        evidence,
        rag_required=True,
        citation_style="footnote",
    )

    assert result["rag_required"] is True
    assert result["citation_style"] == "footnote"
    assert result["reference_note"].endswith("[S1]")
    assert result["verification_report"]["citation_style"] == "footnote"


def test_resource_generation_route_forwards_rag_options(monkeypatch):
    from src.routes import resource_generation

    captured = {}

    class DummySpark:
        def is_configured(self):
            return True

    class DummyCoordinator:
        def process(self, task):
            captured.update(task)
            return {"ok": True}

    monkeypatch.setattr("src.services.spark_service.spark_service", DummySpark())
    monkeypatch.setattr(resource_generation, "_coordinator", DummyCoordinator())

    app = Flask(__name__)
    app.secret_key = "test"
    with app.test_request_context(json={
        "resource_type": "mindmap",
        "topic": "变量",
        "course_id": 7,
        "chapter_ids": [3],
        "rag_required": True,
        "citation_style": "footnote",
        "student_profile": {"cognitive_style": "mixed"},
    }):
        from flask import session

        session["user_id"] = 1
        session["user_role"] = "teacher"
        response, status = resource_generation.generate_single_resource()

    assert status == 200
    assert captured["rag_required"] is True
    assert captured["citation_style"] == "footnote"
    assert captured["options"]["rag_required"] is True
    assert captured["options"]["citation_style"] == "footnote"
    assert captured["course_id"] == 7
    assert captured["chapter_ids"] == [3]


def test_knowledge_graph_import_forwards_rag_options(monkeypatch):
    from src.routes import knowledge_graph_routes

    captured = {}

    def fake_import(course_id, input_type, content=None, file_id=None, rag_required=False, citation_style="bracket", user_id=None, user_role=None):
        captured.update({
            "course_id": course_id,
            "input_type": input_type,
            "content": content,
            "rag_required": rag_required,
            "citation_style": citation_style,
        })
        return {
            "graph_id": "kg_test",
            "rag_required": rag_required,
            "citation_style": citation_style,
            "citation_reliability": {"status": "passed"},
        }

    monkeypatch.setattr(knowledge_graph_routes.syllabus_graph_service, "import_syllabus", fake_import)

    app = Flask(__name__)
    app.secret_key = "test"
    with app.test_request_context(json={
        "input_type": "text",
        "content": "第一章 变量\n变量用于保存数据。",
        "rag_required": True,
        "citation_style": "inline",
    }):
        from flask import session

        session["user_id"] = 1
        session["user_role"] = "teacher"
        response, status = knowledge_graph_routes.import_syllabus(42)

    assert status == 200
    assert captured["rag_required"] is True
    assert captured["citation_style"] == "inline"
    assert response.get_json()["citation_reliability"]["status"] == "passed"


@pytest.fixture
def isolated_graph_app(tmp_path):
    from src.models.course import Course
    from src.models.user import User, db

    app = Flask(__name__)
    app.config.update(
        SQLALCHEMY_DATABASE_URI=f"sqlite:///{tmp_path / 'test.db'}",
        SQLALCHEMY_TRACK_MODIFICATIONS=False,
        TESTING=True,
    )
    db.init_app(app)
    with app.app_context():
        db.create_all()
        user = User(
            username="teacher",
            email="teacher@example.com",
            password_hash="test",
            role="teacher",
        )
        db.session.add(user)
        db.session.commit()
        course = Course(title="Python", description="Python course", teacher_id=user.id)
        db.session.add(course)
        db.session.commit()
        yield app, course
        db.session.remove()
        db.drop_all()


def test_docx_knowledge_graph_import_has_citation_reliability(isolated_graph_app, monkeypatch):
    pytest.importorskip("docx")
    from docx import Document
    from src.services.syllabus_graph_service import syllabus_graph_service

    app, sample_course = isolated_graph_app
    monkeypatch.setattr(syllabus_graph_service, "_parse_text_syllabus", lambda text, course, **kwargs: {
        "course": {"title": course.title, "description": course.description},
        "objectives": [],
        "prerequisites": [],
        "references": [],
        "chapters": [{
            "title": "Python 基础",
            "description": text,
            "knowledge_points": [
                {"title": "变量", "description": "变量用于保存程序运行中的数据。"},
                {"title": "循环语句", "description": "循环语句用于重复执行代码。"},
            ],
        }],
        "raw_text": text,
    })

    doc = Document()
    doc.add_heading("Python 基础", level=1)
    doc.add_paragraph("变量用于保存程序运行中的数据。循环语句用于重复执行代码。")
    stream = io.BytesIO()
    doc.save(stream)
    content = base64.b64encode(stream.getvalue()).decode("ascii")

    with app.app_context():
        result = syllabus_graph_service.import_syllabus(
            sample_course.id,
            "docx",
            content=content,
            rag_required=True,
            citation_style="inline",
        )

    assert result["rag_required"] is True
    assert result["citation_style"] == "inline"
    assert result["source_chunks"]
    assert result["quality_report"]["missing_citation_count"] == 0
    assert result["citation_reliability"]["degradation"] == "unsupported_claims"
    assert result["citation_reliability"]["unsupported_claims"]


def test_pdf_knowledge_graph_import_has_citation_reliability(isolated_graph_app, monkeypatch):
    pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas
    from src.services.syllabus_graph_service import syllabus_graph_service

    app, sample_course = isolated_graph_app
    monkeypatch.setattr(syllabus_graph_service, "_parse_text_syllabus", lambda text, course, **kwargs: {
        "course": {"title": course.title, "description": course.description},
        "objectives": [],
        "prerequisites": [],
        "references": [],
        "chapters": [{
            "title": "Python Basics",
            "description": text,
            "knowledge_points": [
                {"title": "Variables", "description": "Variables store runtime data."},
                {"title": "Loops", "description": "Loops repeat code execution."},
            ],
        }],
        "raw_text": text,
    })

    stream = io.BytesIO()
    pdf = canvas.Canvas(stream)
    pdf.drawString(72, 760, "Python Basics")
    pdf.drawString(72, 740, "Variables store runtime data. Loops repeat code execution.")
    pdf.save()
    content = base64.b64encode(stream.getvalue()).decode("ascii")

    with app.app_context():
        result = syllabus_graph_service.import_syllabus(
            sample_course.id,
            "pdf",
            content=content,
            rag_required=True,
            citation_style="bracket",
        )

    assert result["rag_required"] is True
    assert result["source_chunks"]
    assert result["quality_report"]["missing_citation_count"] == 0
    assert result["citation_reliability"]["degradation"] is None


def test_pdf_parse_merges_rule_fallback_when_llm_under_extracts(monkeypatch):
    from src.services.syllabus_graph_service import syllabus_graph_service

    class Course:
        title = "Java"
        description = "Java course"

    extracted_text = "\n".join([
        "Java 入门",
        "Java 简介",
        "Java 开发环境配置",
        "Java 基础语法",
        "Java 对象和类",
        "Java 继承",
        "Java 集合",
        "Java ArrayList",
        "Java HashMap",
        "Java 多线程",
        "Java 网络编程",
        "Java 高级语法",
        "Java 泛型",
        "Java 8新特性",
    ])

    monkeypatch.setattr(syllabus_graph_service, "_read_pdf_text", lambda content, file_id: extracted_text)
    monkeypatch.setattr(syllabus_graph_service, "_llm_parse_single", lambda text, course, **kwargs: {
        "course": {"title": course.title, "description": course.description},
        "objectives": [],
        "prerequisites": [],
        "references": [],
        "chapters": [{
            "title": "Java 入门",
            "description": "LLM only found one section",
            "knowledge_points": [{"title": "Java 简介", "description": "Java overview"}],
        }],
        "raw_text": text,
    })

    result = syllabus_graph_service.parse_syllabus("pdf", content=b"%PDF fake", course=Course())

    chapter_titles = [chapter["title"] for chapter in result["chapters"]]
    kp_titles = [
        kp["title"]
        for chapter in result["chapters"]
        for kp in chapter.get("knowledge_points", [])
    ]

    assert "Java 入门" in chapter_titles
    assert "Java 集合" in chapter_titles
    assert "Java 多线程" in chapter_titles
    assert "Java 开发环境配置" in kp_titles
    assert "Java HashMap" in kp_titles
    assert "Java 泛型" in kp_titles
    assert len(kp_titles) > 1
