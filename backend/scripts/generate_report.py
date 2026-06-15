"""Generate the project research report as a Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
font.color.rgb = RGBColor(0x00, 0x00, 0x00)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

def add_title(text, level=1):
    """Add a formatted heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体' if level <= 2 else '宋体'
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return heading

def add_para(text, bold=False, indent=True):
    """Add a paragraph."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    if bold:
        run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p

def add_bullet(text, level=0):
    """Add bullet point."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run('• ' + text)
    run.font.size = Pt(11)
    run.font.name = '宋体'
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p

def add_table(headers, data, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = '黑体'
        # Shade header
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): 'D5E8F0'
        })
        shading.append(shading_elem)
    
    # Data rows
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_data))
            run.font.size = Pt(11)
            run.font.name = '宋体'
    
    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = width
    
    return table


# ============================================================
# COVER PAGE
# ============================================================
# Add blank paragraphs for spacing
for _ in range(6):
    doc.add_paragraph()

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('智教星——基于 Spark4.0 大模型\n与多智能体协作的智能教学平台')
run.font.name = '黑体'
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

doc.add_paragraph()  # spacing

# Subtitle
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run('项目研究报告')
run.font.name = '宋体'
run.font.size = Pt(18)
run.bold = True
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()
doc.add_paragraph()

# Info block
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_p.add_run('CRAIC 2026 参赛项目\n\n2026 年 5 月')
run.font.name = '宋体'
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_title('目  录', level=1)
toc_items = [
    '一、项目题目',
    '二、项目摘要',
    '三、项目背景与国内外研究现状',
    '    3.1 项目背景',
    '    3.2 国内外研究现状',
    '    3.3 现有研究不足',
    '四、项目研究内容与技术路线',
    '    4.1 研究内容',
    '    4.2 技术路线',
    '    4.3 系统架构',
    '五、项目创新点',
    '    5.1 教育场景中的多智能体协作系统',
    '    5.2 错题智能归因与靶向练习系统',
    '    5.3 Spark4.0 大模型全流程赋能',
    '    5.4 多维度学习预警系统',
    '    5.5 异步多队列任务架构',
    '六、项目应用前景与社会价值',
    '    6.1 应用领域',
    '    6.2 市场前景',
    '    6.3 社会价值',
    '    6.4 经济效益预估',
    '七、项目存在的问题及改进方向',
    '    7.1 技术难点与不足',
    '    7.2 未解决的问题',
    '    7.3 未来研究方向',
    '    7.4 拓展应用场景',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(item)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

doc.add_page_break()

# ============================================================
# SECTION 1: 项目题目
# ============================================================
add_title('一、项目题目', level=1)
add_para('智教星——基于 Spark4.0 大模型与多智能体协作的智能教学平台研究与应用', bold=True, indent=True)

doc.add_page_break()

# ============================================================
# SECTION 2: 项目摘要
# ============================================================
add_title('二、项目摘要', level=1)
add_para(
    '随着人工智能技术的快速发展，教育信息化正从"数字化"向"智能化"转型。本项目"智教星"是一款基于 '
    'Spark4.0 Ultra 星火大模型的 AI 驱动智能教学管理平台，旨在为教师、学生和管理员提供全方位的智能化教学'
    '解决方案。系统采用前后端分离架构，前端使用 React 19 + Vite 6 + Tailwind CSS，后端使用 Flask 3.0 + '
    'SQLAlchemy + Celery + Redis 构建。核心创新包括：（1）引入多智能体协作系统，通过 Coordinator Agent 协调 '
    'Profile、Exercise、Document、Media、Recommendation 五个专业智能体并行工作，实现学习资源的自动化生成'
    '与一致性保障；（2）构建基于知识图谱的个性化学习路径推荐，结合学生画像模型和学习行为数据，实现因材施教；'
    '（3）开发错题智能归因分析系统，利用 NLP 技术自动识别错题的错误类型（概念理解偏差、计算失误、审题不清等），'
    '并结合 Jaccard 相似度算法生成靶向练习。此外，系统还实现了基于 WebSocket 的实时师生互动、多维度学习数据'
    '可视化与预警等功能。本项目不仅为教育数字化转型提供了可落地的技术解决方案，也为大语言模型在教育场景中的'
    '深度应用提供了实践参考。'
)

doc.add_page_break()

# ============================================================
# SECTION 3: 项目背景与国内外研究现状
# ============================================================
add_title('三、项目背景与国内外研究现状', level=1)

add_title('3.1 项目背景', level=2)
add_para(
    '教育信息化 2.0 行动计划提出，到 2035 年实现教育现代化的战略目标，教育信息化是必经之路。'
    '当前，传统教学平台普遍存在以下问题：'
)
add_bullet('内容生成依赖人工：出题、备课、教案编写耗时耗力，教师工作负担重')
add_bullet('学习路径单一化：固定课程结构无法满足个性化学习需求')
add_bullet('数据分析浅层化：仅做统计展示，缺乏深度的学情诊断和预警')
add_bullet('师生互动效率低：缺乏实时反馈机制，答疑响应慢')
add_bullet('错题管理碎片化：错题整理依赖学生自主完成，缺乏系统性归因分析')

add_title('3.2 国内外研究现状', level=2)

add_para('（1）AI 在教育中的应用', bold=True, indent=False)
add_para(
    '国际上，Knewton 自适应学习平台（2008）率先将知识图谱与自适应算法结合；可汗学院（Khan Academy）'
    '利用 AI 分析学习行为，推荐个性化内容；Coursera 引入 AI 助教，提供 7×24 小时答疑服务。国内，'
    '科大讯飞的智慧课堂系统、百度文心一言教育版、腾讯混元教育解决方案等相继涌现，但多数仍停留在'
    '"问答+内容生成"的浅层应用阶段。'
)

add_para('（2）大语言模型与教育', bold=True, indent=False)
add_para(
    'ChatGPT（2022）引发了教育 AI 的新浪潮，但其在教育场景中的直接应用存在幻觉、知识时效性、'
    '缺乏学科针对性等问题。近年研究指出，基于提示词工程（Prompt Engineering）和检索增强生成（RAG）'
    '的方法能有效提升大模型在教育领域的表现（Liu et al., 2023）。'
)

add_para('（3）多智能体系统', bold=True, indent=False)
add_para(
    '多智能体协作（Multi-Agent System）在软件工程、数据分析等领域已有成功应用。AutoGen（Microsoft, 2023）'
    '证明了多智能体协作在复杂任务上的优势，但在教育领域的系统化应用仍属空白。'
)

add_para('（4）学习分析与预警', bold=True, indent=False)
add_para(
    '学习分析（Learning Analytics）研究已发展十余年，但多数系统仍基于传统的统计分析方法，'
    '缺乏对 AI 生成数据的深度挖掘。基于机器学习的预警模型（如随机森林、XGBoost）在辍学预测上'
    '表现优异（约 80% 准确率），但可解释性不足。'
)

add_title('3.3 现有研究不足', level=2)
add_table(
    ['方面', '不足'],
    [
        ['AI 应用深度', '多为浅层问答，缺乏系统级整合'],
        ['个性化推荐', '基于规则推荐，缺乏动态适配'],
        ['多智能体', '在教育场景的应用仍属空白'],
        ['数据驱动', '预警模型可解释性差，教师难以采纳'],
        ['用户体验', '功能堆砌，缺乏以教学流程为核心的设计'],
    ]
)

doc.add_page_break()

# ============================================================
# SECTION 4: 项目研究内容与技术路线
# ============================================================
add_title('四、项目研究内容与技术路线', level=1)

add_title('4.1 研究内容', level=2)

add_para('（1）多智能体协作系统的构建与优化', bold=True, indent=False)
add_para(
    '设计并实现一个由 5 个专业智能体（Profile Agent、Exercise Agent、Document Agent、'
    'Media Agent、Recommendation Agent）和 1 个协调智能体（Coordinator Agent）组成的多智能体系统。'
    '各智能体通过共享状态（Shared State）和消息总线（Message Bus）进行通信，实现并行工作与一致性保障。'
)
add_bullet('基于 Prompt Engineering 的智能体角色定义')
add_bullet('共享状态管理机制（线程安全的读写控制）')
add_bullet('任务分解与分派策略')
add_bullet('多智能体输出一致性验证')

add_para('（2）基于知识图谱的个性化学习路径推荐', bold=True, indent=False)
add_para(
    '构建包含知识点关联、难度层次、学生掌握度的三维知识图谱，结合学生画像模型，'
    '实现个性化学习路径的自动生成。'
)

add_para('（3）错题智能归因与靶向练习生成', bold=True, indent=False)
add_para(
    '开发自动错题分析系统，利用 NLP 技术将错题归因为四大类型：概念理解偏差、计算失误、'
    '审题不清、其他。基于 Jaccard 相似度算法，生成针对性练习。'
)

add_para('（4）基于 Celery + Redis 的异步任务处理系统', bold=True, indent=False)
add_para(
    '设计多队列异步任务架构，将 AI 生成、邮件发送、数据导出、定时维护等任务分派到不同队列，'
    '提升系统吞吐量。队列设计包括 ai 队列（AI 内容生成）、email 队列、export 队列和 maintenance 队列。'
)

add_para('（5）实时师生互动系统', bold=True, indent=False)
add_para(
    '基于 Flask-SocketIO 和 WebSocket 协议，实现学生举手、实时问答、讨论区等功能的零延迟通信，'
    '延迟低于 2 秒。'
)

add_title('4.2 技术路线', level=2)
add_para('本项目按五个阶段有序推进：', indent=True)
add_bullet('第一阶段：需求分析与技术选型——需求调研、技术栈确定、架构设计')
add_bullet('第二阶段：核心服务开发——Spark4.0 集成、多智能体系统、知识图谱构建')
add_bullet('第三阶段：应用功能开发——AI 出题/答疑、错题分析、学习路径、数据可视化')
add_bullet('第四阶段：性能优化与测试——异步任务、缓存策略、前端优化、测试覆盖')
add_bullet('第五阶段：部署与评估——生产部署、用户测试、效果评估、迭代优化')

add_title('4.3 系统架构', level=2)
add_para(
    '系统采用三层架构设计：用户层（管理员端、教师端、学生端）、API 层（Flask 3.0 提供 RESTful API '
    '和 WebSocket 服务）、服务层（Spark4.0 大模型、多智能体系统、知识图谱构建、学习分析服务）、数据层'
    '（PostgreSQL 数据库、Redis 缓存/消息队列、Elasticsearch 搜索引擎）。'
)

doc.add_page_break()

# ============================================================
# SECTION 5: 项目创新点
# ============================================================
add_title('五、项目创新点', level=1)

add_title('5.1 教育场景中的多智能体协作系统', level=2)
add_para(
    '现有 AI 教育应用多采用单一模型处理所有任务，缺乏专业化和协作性。本项目首创性地将 5 个专业智能体'
    '（习题设计、文档生成、多媒体、资源推荐、项目设计）与协调智能体结合，通过共享状态和消息总线实现'
    '异步通信，引入一致性验证机制确保多智能体输出在知识点覆盖和难度上的一致性。各智能体可独立工作，'
    '系统具有高可扩展性和容错能力。'
)

add_title('5.2 错题智能归因与靶向练习系统', level=2)
add_para(
    '传统错题本仅记录错题，缺乏自动归因分析和针对性练习生成。本项目基于 NLP 关键词匹配算法，'
    '将错题自动归因为四大类型（概念理解偏差、计算失误、审题不清、其他），使用 Jaccard 相似度算法'
    '生成靶向练习，结合 Spark4.0 大模型生成错因分析和改进建议，实现错题掌握度追踪'
    '（unmastered → reviewing → mastered）。'
)

add_title('5.3 基于 Spark4.0 大模型的 AI 全流程赋能', level=2)
add_para('当前 AI 应用多集中在单一环节，缺乏系统性整合。本项目实现：')
add_bullet('备课环节：AI 自动生成学习目标、知识要点、代码示例、练习题、答案解析')
add_bullet('教学环节：AI 实时答疑、智能推荐、学习路径规划')
add_bullet('评价环节：AI 出题、自动评测、个性化反馈')
add_bullet('分析环节：AI 生成学习报告、预警识别、知识图谱构建')
add_bullet('管理环节：AI 辅助排课、资源推荐、质量评估')

add_title('5.4 多维度学习预警系统', level=2)
add_para(
    '传统预警基于单一指标，准确率和可解释性不足。本项目融合多源数据（学习进度、错题数量、'
    '编程提交、活跃度）计算风险评分，支持多维度预警类型（不活跃、成绩下滑、低分预警、知识缺口），'
    '实现分级预警机制（info → warning → critical）。'
)

add_title('5.5 异步多队列任务架构', level=2)
add_para(
    'AI 调用同步阻塞会导致请求堆积和响应超时。本项目基于 Celery + Redis 构建 6 个专用队列'
    '（default、ai、email、export、maintenance、high_priority），支持任务超时控制、重试机制、'
    '进度报告，通过异步 AI 生成任务，前端通过 SSE/WebSocket 获取实时进度。'
)

doc.add_page_break()

# ============================================================
# SECTION 6: 项目应用前景与社会价值
# ============================================================
add_title('六、项目应用前景与社会价值', level=1)

add_title('6.1 应用领域', level=2)
add_table(
    ['领域', '应用场景', '价值'],
    [
        ['K12 教育', '智慧课堂、课后辅导', '减轻教师负担，提升教学效率'],
        ['高等教育', '在线课程、编程教学', '支持大规模个性化教学'],
        ['职业培训', '技能培训、认证考试', '自动化出题、学习跟踪'],
        ['企业培训', '员工能力评估、在线学习', '降低培训成本，提升效果'],
    ],
    [Cm(3), Cm(4.5), Cm(4.5)]
)

add_title('6.2 市场前景', level=2)
add_bullet('中国在线教育市场规模预计 2025 年达 5000 亿元')
add_bullet('AI + 教育细分市场年增长率超过 30%')
add_bullet('本平台可作为 SaaS 服务，面向学校、教育机构提供订阅服务')

add_title('6.3 社会价值', level=2)
add_bullet('教育资源均衡：通过 AI 能力弥补优质教师不足，使偏远地区学生也能享受优质教育')
add_bullet('因材施教：个性化学习路径让每个学生都能按照自己的节奏学习')
add_bullet('教师减负：自动化出题、批改、统计等功能，让教师将更多精力投入到教学创新上')
add_bullet('终身学习：错题本和学习轨迹追踪功能，支持学生长期学习成长')
add_bullet('数据驱动决策：多维度数据分析，为学校和教育管理者提供科学决策支持')

add_title('6.4 经济效益预估', level=2)
add_table(
    ['指标', '预估值'],
    [
        ['教师备课时间节省', '40-60%'],
        ['学生学习效率提升', '30-50%'],
        ['平台运维成本降低', '30%（相比传统方案）'],
        ['单校年运营成本节省', '约 50-100 万元'],
    ]
)

doc.add_page_break()

# ============================================================
# SECTION 7: 项目存在的问题及改进方向
# ============================================================
add_title('七、项目存在的问题及改进方向', level=1)

add_title('7.1 技术难点与不足', level=2)

add_para('（1）AI 调用缺乏熔断和降级机制', bold=True, indent=False)
add_para('问题：Spark4.0 API 超时或不可用时，请求会堆积导致服务不可用。')
add_para('改进：引入 Circuit Breaker 模式，设置超时阈值和熔断阈值，API 不可用时降级为缓存或模板回答。')

add_para('（2）数据库缺少索引优化', bold=True, indent=False)
add_para('问题：大量查询操作缺乏数据库索引，性能低下。')
add_para('改进：为常用查询字段添加索引（user_id + course_id、behavior_type + created_at 等）。')

add_para('（3）多智能体系统缺乏动态协作能力', bold=True, indent=False)
add_para('问题：当前智能体协作是静态的，无法根据任务难度和类型动态调整 Agent 组合。')
add_para('改进：引入强化学习算法，让 Coordinator Agent 学习最优的 Agent 调度策略。')

add_para('（4）知识图谱构建依赖人工标注', bold=True, indent=False)
add_para('问题：知识点标签和关联关系仍需人工定义，缺乏自动发现和构建能力。')
add_para('改进：利用 NLP 技术自动提取教材中的知识点及其关联，构建动态知识图谱。')

add_para('（5）前端状态管理分散', bold=True, indent=False)
add_para('问题：用户状态和课程数据在多个组件中独立管理，跨组件通信困难。')
add_para('改进：引入 Zustand 或 Redux 等全局状态管理方案。')

add_title('7.2 未解决的问题', level=2)
add_table(
    ['问题', '说明', '改进方向'],
    [
        ['多模态学习', '当前仅支持文本交互', '引入图像识别（拍照搜题）、语音交互'],
        ['向量数据库', '缺乏长期记忆和语义搜索', '引入 Milvus/Chroma 存储对话历史'],
        ['在线监考', '考试缺乏防作弊机制', '引入人脸识别、屏幕监控'],
        ['跨平台支持', '移动端体验不佳', '开发独立 App 或优化 PWA'],
        ['国际化', '仅支持中文', '多语言支持'],
    ],
    [Cm(3), Cm(4), Cm(5)]
)

add_title('7.3 未来研究方向', level=2)
add_bullet('RAG 增强：结合教材、课程标准等外部知识源，提升 AI 生成内容的准确性和针对性')
add_bullet('联邦学习：在保护学生隐私的前提下，利用多校数据训练更精准的推荐模型')
add_bullet('情感计算：通过学习行为分析学生情绪状态，提供及时的情感支持和鼓励')
add_bullet('数字孪生：构建学生数字孪生模型，模拟不同学习策略的效果，辅助教学决策')
add_bullet('边缘计算：将部分 AI 推理能力部署到边缘设备，降低网络依赖和延迟')

add_title('7.4 拓展应用场景', level=2)
add_bullet('AI 助教：作为教师的智能助手，辅助备课、批改、答疑')
add_bullet('教育大模型：针对教育场景微调专用大模型，替代通用模型')
add_bullet('教育数据分析平台：面向教育管理部门提供区域级学情分析和决策支持')
add_bullet('开放 API：将 AI 能力封装为 API，供第三方教育应用调用')

# ============================================================
# Save
# ============================================================
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    '智教星项目研究报告.docx'
)
doc.save(output_path)
print(f'文档已生成: {output_path}')
