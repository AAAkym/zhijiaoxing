import base64
import io
import json
import logging
import re
import time
import uuid
import zipfile
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from collections import defaultdict
from xml.etree import ElementTree as ET

from docx import Document
from sqlalchemy.exc import OperationalError

from src.models.course import Course
from src.models.knowledge_base import (
    CourseChapter,
    CourseSyllabus,
    KnowledgeGraphEdge,
    KnowledgeGraphNode,
    KnowledgePoint,
    KnowledgeSourceChunk,
)
from src.models.user import db
from src.services.multi_agent.knowledge_graph_agent import knowledge_graph_agent
from src.services.rag_citation_service import rag_citation_service

logger = logging.getLogger(__name__)


NODE_TYPES = {
    "course", "chapter", "knowledge_point", "objective", "skill",
    "case", "exercise", "resource",
}
EDGE_TYPES = {
    "contains", "prerequisite", "related", "supports_objective",
    "applies_to", "assesses", "recommended_after",
}


class SyllabusGraphService:
    _graph_write_lock = threading.RLock()
    _node_cache = {}
    _edge_cache = {}

    def _reset_node_cache(self):
        self._node_cache.clear()
        self._edge_cache.clear()

    def _cached_upsert_node(self, graph_id, course_id, node_type, label, description="", category=None, source_chunk_ids=None, properties=None):
        """带内存缓存的节点 upsert，避免清空图谱后重复 DB 查询"""
        node_type = node_type if node_type in NODE_TYPES else "knowledge_point"
        label = (label or "").strip()[:200]
        cache_key = (course_id, node_type, label)
        node = self._node_cache.get(cache_key)
        is_new = node is None
        if not node:
            node = KnowledgeGraphNode(course_id=course_id, node_type=node_type, label=label, graph_id=graph_id)
            db.session.add(node)
            db.session.flush()
            self._node_cache[cache_key] = node
        node.description = description or node.description
        node.category = category or node.category
        node.source_chunk_ids = json.dumps(source_chunk_ids or [], ensure_ascii=False)
        node.properties = json.dumps(properties or {}, ensure_ascii=False)
        return node, is_new

    def clear_graph(self, course_id):
        """Remove all knowledge graph data for a course."""
        course = Course.query.get(course_id)
        if not course:
            return {"error": "Course not found"}
        node_count = KnowledgeGraphNode.query.filter_by(course_id=course_id).count()
        edge_count = KnowledgeGraphEdge.query.filter_by(course_id=course_id).count()
        chunk_count = KnowledgeSourceChunk.query.filter_by(course_id=course_id).count()
        with self._graph_write_lock:
            self._run_with_sqlite_lock_retry(
                lambda: self._clear_existing_graph(course_id, delete_all_chunks=True),
                operation="clear graph",
            )
            self._commit_with_sqlite_lock_retry(operation="clear graph commit")
        return {
            "deleted_nodes": node_count,
            "deleted_edges": edge_count,
            "deleted_source_chunks": chunk_count,
            "message": f"已清除课程 {course.title} 的知识图谱",
        }

    def import_syllabus(
        self,
        course_id,
        input_type,
        content=None,
        file_id=None,
        rag_required=False,
        citation_style="bracket",
        user_id=None,
        user_role=None,
    ):
        course = Course.query.get(course_id)
        if not course:
            return {"error": "Course not found"}

        perf = {}

        t0 = time.perf_counter()
        normalized = self.parse_syllabus(
            input_type,
            content=content,
            file_id=file_id,
            course=course,
            user_id=user_id,
            user_role=user_role,
        )
        perf["parse_ms"] = round((time.perf_counter() - t0) * 1000, 2)

        graph_id = f"kg_{course_id}_{uuid.uuid4().hex[:8]}"

        t2 = time.perf_counter()
        with self._graph_write_lock:
            self._run_with_sqlite_lock_retry(
                lambda: self._clear_existing_graph(course_id),
                operation="clear existing graph before import",
            )
        self._reset_node_cache()
        source_chunks = self._create_source_chunks(course_id, normalized)
        node_cache = {}
        created_nodes = 0
        created_edges = 0
        empty_kp_ids = []

        course_node, is_new = self._cached_upsert_node(
            graph_id, course_id, "course", normalized["course"].get("title") or course.title,
            normalized["course"].get("description") or course.description or "",
            category="课程", source_chunk_ids=[c.id for c in source_chunks[:1]],
        )
        node_cache[("course", course_node.label)] = course_node
        created_nodes += int(is_new)

        objective_nodes = []
        for objective in normalized.get("objectives", []):
            node, is_new = self._cached_upsert_node(
                graph_id, course_id, "objective", objective, objective,
                category="目标", source_chunk_ids=self._chunk_ids_for_text(source_chunks, objective),
            )
            objective_nodes.append(node)
            created_nodes += int(is_new)
            created_edges += self._upsert_edge(
                graph_id, course_id, course_node.id, node.id, "supports_objective", 1.0, 0.95,
                self._chunk_ids_for_text(source_chunks, objective),
            )

        previous_chapter_node = None
        kp_nodes = []
        for chapter_index, chapter in enumerate(normalized.get("chapters", []), start=1):
            ch_label = chapter.get("title") or f"第{chapter_index}章"
            ch_node, is_new = self._cached_upsert_node(
                graph_id, course_id, "chapter", ch_label, chapter.get("description", ""),
                category=chapter.get("chapter_type", "theory"),
                source_chunk_ids=self._chunk_ids_for_text(source_chunks, ch_label),
                properties={
                    "order_index": chapter_index,
                    "teaching_hours": chapter.get("teaching_hours", 0),
                    "main_content": self._build_chapter_main_content(chapter),
                    "annotation": self._build_chapter_annotation(chapter),
                    "key_points": [
                        (kp.get("title") if isinstance(kp, dict) else str(kp))
                        for kp in chapter.get("knowledge_points", [])[:8]
                    ],
                },
            )
            created_nodes += int(is_new)
            created_edges += self._upsert_edge(
                graph_id, course_id, course_node.id, ch_node.id, "contains", 1.0, 1.0,
                self._chunk_ids_for_text(source_chunks, ch_label),
            )
            if previous_chapter_node:
                created_edges += self._upsert_edge(
                    graph_id, course_id, previous_chapter_node.id, ch_node.id, "recommended_after", 0.7, 0.8, []
                )
            previous_chapter_node = ch_node

            for objective_node in objective_nodes:
                created_edges += self._upsert_edge(
                    graph_id, course_id, ch_node.id, objective_node.id, "supports_objective", 0.7, 0.75, []
                )

            previous_kp_node = None
            for kp_index, kp in enumerate(chapter.get("knowledge_points", []), start=1):
                kp_label = kp.get("title") if isinstance(kp, dict) else str(kp)
                if not kp_label:
                    continue
                kp_desc = kp.get("description", "") if isinstance(kp, dict) else ""
                kp_desc = self._build_kp_description(kp_label, kp_desc, ch_label)
                kp_props = {"chapter": ch_label, "order_index": kp_index}
                # 为知识点生成内容摘要summary
                if kp_desc:
                    kp_props["summary"] = kp_desc[:80]
                elif isinstance(kp, dict):
                    # 从相关概念和内容中生成摘要
                    related = kp.get("related_concepts", [])
                    if related:
                        kp_props["summary"] = "相关：" + "、".join(str(r) for r in related[:3])
                if isinstance(kp, dict):
                    if kp.get("difficulty"):
                        kp_props["difficulty"] = kp["difficulty"]
                    if kp.get("assessed_by"):
                        kp_props["assessed_by"] = str(kp["assessed_by"])
                kp_node, is_new = self._cached_upsert_node(
                    graph_id, course_id, "knowledge_point", kp_label,
                    kp_desc,
                    category=kp.get("category", "核心知识点") if isinstance(kp, dict) else "核心知识点",
                    source_chunk_ids=self._chunk_ids_for_text(source_chunks, kp_label),
                    properties=kp_props,
                )
                if not (kp_desc or "").strip():
                    empty_kp_ids.append(kp_node.id)
                kp_nodes.append(kp_node)
                created_nodes += int(is_new)
                created_edges += self._upsert_edge(
                    graph_id, course_id, ch_node.id, kp_node.id, "contains", 1.0, 0.95,
                    self._chunk_ids_for_text(source_chunks, kp_label),
                )
                if previous_kp_node:
                    created_edges += self._upsert_edge(
                        graph_id, course_id, previous_kp_node.id, kp_node.id, "recommended_after", 0.65, 0.75, []
                    )
                previous_kp_node = kp_node

                for related in (kp.get("related_concepts", []) if isinstance(kp, dict) else []):
                    related_node, is_new = self._cached_upsert_node(
                        graph_id, course_id, "knowledge_point", str(related), "", "关联概念",
                        self._chunk_ids_for_text(source_chunks, str(related)),
                    )
                    empty_kp_ids.append(related_node.id)
                    created_nodes += int(is_new)
                    created_edges += self._upsert_edge(
                        graph_id, course_id, kp_node.id, related_node.id, "related", 0.6, 0.7,
                        self._chunk_ids_for_text(source_chunks, str(related)),
                    )

        for prereq in normalized.get("prerequisites", []):
            prereq_node, is_new = self._cached_upsert_node(
                graph_id, course_id, "skill", str(prereq), "", "先修技能",
                self._chunk_ids_for_text(source_chunks, str(prereq)),
            )
            created_nodes += int(is_new)
            for kp_node in kp_nodes[:12]:
                created_edges += self._upsert_edge(
                    graph_id, course_id, prereq_node.id, kp_node.id, "prerequisite", 0.9, 0.85,
                    self._chunk_ids_for_text(source_chunks, str(prereq)),
                )

        # 应用AI提取的跨章节关系
        cross_relations = getattr(self, '_cross_chapter_relations', None) or []
        if cross_relations:
            created_edges += self._apply_agent_relations(graph_id, course_id, kp_nodes, cross_relations)

        self._commit_with_sqlite_lock_retry(operation="import syllabus commit")
        perf["build_ms"] = round((time.perf_counter() - t2) * 1000, 2)

        # 清理没有核心内容的知识点节点及其关联边（使用创建时记录的ID，避免全表扫描）
        if empty_kp_ids:
            KnowledgeGraphEdge.query.filter_by(course_id=course_id).filter(
                (KnowledgeGraphEdge.source_node_id.in_(empty_kp_ids)) |
                (KnowledgeGraphEdge.target_node_id.in_(empty_kp_ids))
            ).delete(synchronize_session=False)
            KnowledgeGraphNode.query.filter(KnowledgeGraphNode.id.in_(empty_kp_ids)).delete(synchronize_session=False)
            self._commit_with_sqlite_lock_retry(operation="cleanup empty nodes commit")

        quality_report = self.build_quality_report(course_id)
        citation_reliability = self._build_graph_citation_reliability(
            course_id,
            rag_required=rag_required,
            citation_style=citation_style,
        )
        return {
            "graph_id": graph_id,
            "nodes_created": created_nodes,
            "edges_created": created_edges,
            "source_chunks": [chunk.to_dict() for chunk in source_chunks],
            "quality_report": quality_report,
            "citation_reliability": citation_reliability,
            "rag_required": bool(rag_required),
            "citation_style": citation_style or "bracket",
            "performance": perf,
        }

    def parse_syllabus(self, input_type, content=None, file_id=None, course=None, user_id=None, user_role=None):
        input_type = self._detect_input_type(input_type, content, file_id)
        if input_type == "docx":
            text = self._read_docx_text(content, file_id)
        elif input_type == "pdf":
            text = self._read_pdf_text(content, file_id)
        elif input_type == "json":
            return self._normalize_structured(content, course)
        elif input_type in ("text", "markdown"):
            text = content or ""
        else:
            raise ValueError("input_type must be one of: json, text, markdown, docx, pdf")
        if not (text or "").strip():
            raise ValueError("解析内容为空，请确认文档或PDF内容不为空")
        analysis = self._parse_text_syllabus(text, course, user_id=user_id, user_role=user_role)
        if not analysis:
            detail = getattr(self, '_last_parse_error', '') or '未知原因'
            raise ValueError(f"AI解析失败，无法从文档中提取知识点。原因：{detail}。请重试或检查文档内容（建议上传结构清晰的课程大纲，文本长度不超过2万字）")
        analysis["_analysis_budget"] = self._build_analysis_budget(text, analysis)
        return self._postprocess_analysis(analysis)

    def _detect_input_type(self, input_type, content=None, file_id=None):
        declared = (input_type or "text").lower()
        if declared == "json":
            return declared
        try:
            binary = self._to_binary(content, file_id)
        except Exception:
            return declared
        if binary.startswith(b"%PDF"):
            return "pdf"
        if binary.startswith((b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")):
            return "docx"
        return declared

    def get_graph(self, course_id, node_type=None, edge_type=None, include_sources=True):
        nodes_query = KnowledgeGraphNode.query.filter_by(course_id=course_id)
        if node_type:
            nodes_query = nodes_query.filter_by(node_type=node_type)
        nodes = nodes_query.order_by(KnowledgeGraphNode.id).all()
        node_ids = {n.id for n in nodes}

        edges_query = KnowledgeGraphEdge.query.filter_by(course_id=course_id)
        if edge_type:
            edges_query = edges_query.filter_by(edge_type=edge_type)
        edges = [
            e for e in edges_query.order_by(KnowledgeGraphEdge.id).all()
            if e.source_node_id in node_ids and e.target_node_id in node_ids
        ]

        source_map = {}
        if include_sources:
            source_map = {c.id: c.to_dict() for c in KnowledgeSourceChunk.query.filter_by(course_id=course_id).all()}

        return {
            "nodes": [self._format_node_for_graph(n, source_map if include_sources else None) for n in nodes],
            "edges": [e.to_dict(include_sources=include_sources) for e in edges],
            "metrics": self._build_metrics(nodes, edges),
            "quality_report": self.build_quality_report(course_id),
        }

    def build_course_profile(self, course_id):
        graph = self.get_graph(course_id, include_sources=False)
        nodes = graph["nodes"]
        edges = graph["edges"]
        kp_count = sum(1 for n in nodes if n["node_type"] == "knowledge_point")
        chapter_count = sum(1 for n in nodes if n["node_type"] == "chapter")
        prereq_count = sum(1 for e in edges if e["edge_type"] == "prerequisite")
        practice_count = sum(1 for n in nodes if n["node_type"] in ("case", "exercise", "skill"))
        return {
            "knowledge_density": round(kp_count / max(chapter_count, 1), 2),
            "core_knowledge_count": kp_count,
            "chapter_count": chapter_count,
            "prerequisite_strength": round(prereq_count / max(kp_count, 1), 2),
            "practice_ratio": round(practice_count / max(len(nodes), 1), 2),
            "difficulty": "advanced" if kp_count > 40 or prereq_count > 20 else "intermediate" if kp_count > 15 else "beginner",
        }

    def build_quality_report(self, course_id):
        nodes = KnowledgeGraphNode.query.filter_by(course_id=course_id).all()
        edges = KnowledgeGraphEdge.query.filter_by(course_id=course_id).all()
        chunk_count = KnowledgeSourceChunk.query.filter_by(course_id=course_id).count()
        connected = set()
        for edge in edges:
            connected.add(edge.source_node_id)
            connected.add(edge.target_node_id)
        missing_citations = sum(1 for n in nodes if not json.loads(n.source_chunk_ids or "[]"))
        return {
            "node_count": len(nodes),
            "edge_count": len(edges),
            "isolated_node_count": sum(1 for n in nodes if n.id not in connected),
            "source_coverage_rate": round((len(nodes) - missing_citations) / max(len(nodes), 1) * 100, 1),
            "average_edge_weight": round(sum(e.weight or 0 for e in edges) / max(len(edges), 1), 2),
            "missing_citation_count": missing_citations,
            "source_chunk_count": chunk_count,
        }

    def _build_graph_citation_reliability(self, course_id, rag_required=False, citation_style="bracket"):
        chunks = KnowledgeSourceChunk.query.filter_by(course_id=course_id).all()
        citations = []
        for chunk in chunks:
            item = chunk.to_dict()
            citations.append({
                "source_id": item.get("reference_code") or item.get("source_id"),
                "source_type": item.get("source_type"),
                "title": item.get("title"),
                "excerpt": item.get("excerpt"),
                "location": item.get("location"),
                "url": item.get("url"),
                "source_chunk_id": item.get("id"),
                "confidence": 0.95,
            })

        node_texts = []
        for node in KnowledgeGraphNode.query.filter_by(course_id=course_id).all():
            refs = []
            for chunk_id in json.loads(node.source_chunk_ids or "[]"):
                chunk = next((c for c in chunks if c.id == chunk_id), None)
                if chunk:
                    refs.append(f"[{chunk.reference_code}]")
            text = " ".join(filter(None, [node.label, node.description, " ".join(refs)]))
            if text:
                node_texts.append(text)

        edge_texts = []
        for edge in KnowledgeGraphEdge.query.filter_by(course_id=course_id).all():
            refs = []
            for chunk_id in json.loads(edge.evidence_chunk_ids or "[]"):
                chunk = next((c for c in chunks if c.id == chunk_id), None)
                if chunk:
                    refs.append(f"[{chunk.reference_code}]")
            text = " ".join(filter(None, [edge.edge_type, " ".join(refs)]))
            if text:
                edge_texts.append(text)

        verification = rag_citation_service.verify(
            "\n".join(node_texts + edge_texts),
            citations,
            rag_required=rag_required,
            citation_style=citation_style,
        )
        return {
            "status": verification["status"],
            "citation_coverage_score": verification["citation_coverage_score"],
            "unsupported_claims": verification["unsupported_claims"],
            "citation_issues": verification["citation_issues"],
            "degradation": verification["degradation"],
            "degraded": verification["degraded"],
        }

    def _normalize_structured(self, content, course):
        data = content
        if isinstance(content, str):
            data = json.loads(content or "{}")
        data = data or {}
        chapters = data.get("chapters") or data.get("units") or []
        normalized_chapters = []
        for idx, ch in enumerate(chapters, start=1):
            if isinstance(ch, str):
                normalized_chapters.append({"title": ch, "knowledge_points": []})
                continue
            key_points = ch.get("knowledge_points") or ch.get("key_points") or ch.get("points") or []
            normalized_chapters.append({
                "title": ch.get("title") or ch.get("name") or f"第{idx}章",
                "description": ch.get("description") or ch.get("summary") or "",
                "teaching_hours": ch.get("teaching_hours") or ch.get("hours") or 0,
                "chapter_type": ch.get("chapter_type") or "theory",
                "knowledge_points": self._normalize_kp_list(key_points),
            })
        return {
            "course": {
                "title": data.get("title") or data.get("course_title") or getattr(course, "title", ""),
                "description": data.get("description") or getattr(course, "description", ""),
            },
            "objectives": self._normalize_text_list(data.get("objectives") or data.get("course_objectives") or []),
            "prerequisites": self._normalize_text_list(data.get("prerequisites") or data.get("prerequisite_courses") or []),
            "references": self._normalize_text_list(data.get("references") or []),
            "chapters": normalized_chapters,
            "raw_text": json.dumps(data, ensure_ascii=False),
        }

    def _parse_text_syllabus(self, text, course, user_id=None, user_role=None):
        """Extract knowledge points with deterministic rules plus optional LLM enrichment.

        优化：规则解析同步执行（快），LLM 富化在守护线程中运行并设硬性等待上限。
        超时后立即返回规则结果，避免 ThreadPoolExecutor 的阻塞式 shutdown 拖死整体。
        """
        text = text or ""

        rule_result = None
        rule_error = None

        # 1) 规则解析：同步执行，通常 1-3s 完成，产出完整结构
        try:
            rule_result = self._rule_based_parse(text, course)
        except Exception as exc:
            rule_error = exc
            logger.warning("规则解析异常: %s", exc)

        # 2) LLM 富化：守护线程执行，join 设硬超时，超时即放弃（非阻塞）
        llm_box = {}
        use_chunked = len(text) > self.LITE_CONTEXT_LIMIT
        if use_chunked:
            logger.info("文本长度 %d 超过上下文上限 %d，启用分块解析", len(text), self.LITE_CONTEXT_LIMIT)

        def _run_llm_parse():
            try:
                if use_chunked:
                    llm_box['result'] = self._llm_based_parse(text, course, user_id=user_id, user_role=user_role)
                else:
                    llm_box['result'] = self._llm_parse_single(text, course, user_id=user_id, user_role=user_role)
            except Exception as exc:
                llm_box['error'] = exc

        llm_thread = threading.Thread(target=_run_llm_parse, daemon=True)
        llm_thread.start()

        # LLM 富化等待上限：超时则用规则结果，不再阻塞等待底层 HTTP 调用结束
        LLM_WAIT_SECONDS = 120
        llm_started = time.perf_counter()
        llm_thread.join(timeout=LLM_WAIT_SECONDS)
        elapsed = time.perf_counter() - llm_started

        llm_result = llm_box.get('result')
        llm_error = llm_box.get('error')

        if llm_result:
            logger.info("AI解析完成，耗时 %.2fs", elapsed)
        elif llm_thread.is_alive():
            logger.warning("AI解析超时（%.0fs 等待上限），使用规则解析结果", elapsed)
            self._last_parse_error = "AI解析超时，已使用规则解析结果"
        elif llm_error:
            logger.warning("AI解析异常（耗时 %.2fs），使用规则解析结果兜底: %s", elapsed, llm_error)
            self._last_parse_error = f"AI解析异常: {llm_error}"

        # 3) 兜底：两者都无结果时抛错
        if rule_result is None and llm_result is None:
            if rule_error:
                raise rule_error
            if llm_error:
                raise llm_error
            return None

        if rule_result and llm_result:
            return self._merge_parse_results(rule_result, llm_result, course)
        return llm_result or rule_result

    def _merge_parse_results(self, rule_result, llm_result, course):
        """Merge rule-based and LLM-based parse results."""
        merged = {
            "course": llm_result.get("course") or rule_result.get("course", {}),
            "raw_text": rule_result.get("raw_text", ""),
            "_analysis_budget": rule_result.get("_analysis_budget") or llm_result.get("_analysis_budget"),
        }

        # 合并目标：LLM 语义理解更准确，优先；规则补充遗漏
        rule_objectives = set(o for o in rule_result.get("objectives", []) if len(o) > 2)
        llm_objectives = set(o for o in llm_result.get("objectives", []) if len(o) > 2)
        merged["objectives"] = list(llm_objectives | rule_objectives)[:15]

        # 合并先修要求
        rule_prereqs = set(p for p in rule_result.get("prerequisites", []) if len(p) > 1)
        llm_prereqs = set(p for p in llm_result.get("prerequisites", []) if len(p) > 1)
        merged["prerequisites"] = list(llm_prereqs | rule_prereqs)[:12]

        merged["references"] = list(dict.fromkeys(
            rule_result.get("references", []) + llm_result.get("references", [])
        ))[:15]

        # 合并章节与知识点（核心逻辑）        # 策略：以 LLM 结果为骨架（结构更合理），用规则结果补充遗漏的知识点
        llm_chapters = llm_result.get("chapters", [])
        rule_chapters = rule_result.get("chapters", [])
        llm_kp_count = self._count_knowledge_points(llm_chapters)
        rule_kp_count = self._count_knowledge_points(rule_chapters)
        llm_chapter_count = len(llm_chapters)
        rule_chapter_count = len(rule_chapters)

        rule_has_richer_structure = (
            rule_chapter_count >= max(3, llm_chapter_count * 2)
            and rule_kp_count >= max(6, llm_kp_count + 2)
        )
        if rule_kp_count >= max(20, llm_kp_count * 2) or rule_has_richer_structure:
            llm_chapters = self._merge_llm_kps_into_rule_chapters(rule_chapters, llm_chapters)

        llm_kp_set = set()
        for ch in llm_chapters:
            for kp in ch.get("knowledge_points", []):
                title = kp.get("title", "") if isinstance(kp, dict) else str(kp)
                if title:
                    llm_kp_set.add(title.strip())

        missing_kps = []
        for ch in rule_chapters:
            for kp in ch.get("knowledge_points", []):
                title = kp.get("title", "") if isinstance(kp, dict) else str(kp)
                title = self._summarize_to_kp_title(title)
                if title and title not in llm_kp_set:
                    missing_kps.append({
                        "title": title,
                        "description": kp.get("description", "") if isinstance(kp, dict) else "",
                        "category": kp.get("category", "核心知识点") if isinstance(kp, dict) else "核心知识点",
                    })

        if missing_kps and llm_chapters:
            self._distribute_missing_kps(llm_chapters, missing_kps)
        elif missing_kps and not llm_chapters:
            llm_chapters.append({
                "title": "其他章节",
                "description": "其他未归类的知识点",
                "knowledge_points": missing_kps[:self._analysis_budget(merged).get("total_kp_limit", 60)],
                "teaching_hours": 0,
                "chapter_type": "theory",
            })

        merged["chapters"] = llm_chapters if llm_chapters else rule_chapters

        rule_kp_desc = {}
        for ch in rule_chapters:
            for kp in ch.get("knowledge_points", []):
                if isinstance(kp, dict):
                    t = kp.get("title", "").strip()
                    d = kp.get("description", "").strip()
                    if t and d:
                        rule_kp_desc[t] = d

        for ch in merged["chapters"]:
            for kp in ch.get("knowledge_points", []):
                if isinstance(kp, dict):
                    t = kp.get("title", "").strip()
                    if t and not kp.get("description", "").strip() and t in rule_kp_desc:
                        kp["description"] = rule_kp_desc[t]

        return self._postprocess_analysis(merged)

    def _build_analysis_budget(self, text, analysis=None):
        text = text or ""
        analysis = analysis or {}
        chapters = [ch for ch in analysis.get("chapters", []) or [] if isinstance(ch, dict)]
        candidate_count = self._count_knowledge_points(chapters)
        line_count = len([line for line in text.splitlines() if line.strip()])
        text_len = len(text)
        chapter_count = max(len(chapters), 1)

        estimated_total = max(
            candidate_count,
            min(260, max(12, text_len // 110, line_count // 3)),
        )
        if text_len < 2500 and candidate_count <= 20:
            estimated_total = max(candidate_count, 8)
        elif text_len > 12000 or line_count > 500:
            estimated_total = max(estimated_total, min(360, candidate_count or line_count // 2))

        per_chapter = max(4, min(28, int((estimated_total / chapter_count) + 3)))
        return {
            "text_length": text_len,
            "line_count": line_count,
            "chapter_count": chapter_count,
            "candidate_kp_count": candidate_count,
            "total_kp_limit": int(max(estimated_total, candidate_count)),
            "per_chapter_kp_limit": per_chapter,
        }

    def _analysis_budget(self, analysis):
        budget = (analysis or {}).get("_analysis_budget") or {}
        if budget:
            return budget
        return self._build_analysis_budget((analysis or {}).get("raw_text", ""), analysis)

    def _chapter_kp_limit(self, chapter, budget):
        current_count = len(chapter.get("knowledge_points", []) or [])
        base = int((budget or {}).get("per_chapter_kp_limit") or 12)
        if current_count > base:
            return min(current_count, max(base, int(current_count * 0.85)))
        return base

    def _count_knowledge_points(self, chapters):
        return sum(len(ch.get("knowledge_points", []) or []) for ch in chapters or [] if isinstance(ch, dict))

    def _merge_llm_kps_into_rule_chapters(self, rule_chapters, llm_chapters):
        merged = [dict(ch) for ch in rule_chapters or [] if isinstance(ch, dict)]
        if not merged:
            return llm_chapters or []

        def norm(value):
            return self._clean_compare_text(value)

        for llm_ch in llm_chapters or []:
            if not isinstance(llm_ch, dict):
                continue
            llm_title = llm_ch.get("title", "")
            target = None
            for rule_ch in merged:
                rule_title = rule_ch.get("title", "")
                if norm(llm_title) and (norm(llm_title) in norm(rule_title) or norm(rule_title) in norm(llm_title)):
                    target = rule_ch
                    break
            if target is None:
                target = merged[0]

            existing = {
                self._clean_compare_text(kp.get("title", "") if isinstance(kp, dict) else str(kp))
                for kp in target.get("knowledge_points", []) or []
            }
            for kp in llm_ch.get("knowledge_points", []) or []:
                title = kp.get("title", "") if isinstance(kp, dict) else str(kp)
                compare = self._clean_compare_text(title)
                if compare and compare not in existing:
                    target.setdefault("knowledge_points", []).append(kp)
                    existing.add(compare)
        return merged

    def _distribute_missing_kps(self, chapters, missing_kps):
        if not chapters:
            return
        chapter_count = len(chapters)
        limit = self._analysis_budget({"chapters": chapters}).get("total_kp_limit", 120)
        for index, kp in enumerate(missing_kps[:limit]):
            target = chapters[index % chapter_count]
            existing = {
                self._clean_compare_text(item.get("title", "") if isinstance(item, dict) else str(item))
                for item in target.get("knowledge_points", []) or []
            }
            compare = self._clean_compare_text(kp.get("title", "") if isinstance(kp, dict) else str(kp))
            if compare and compare not in existing:
                target.setdefault("knowledge_points", []).append(kp)

    def _postprocess_analysis(self, analysis):
        analysis = analysis or {}
        chapters = analysis.get("chapters") or []
        budget = self._analysis_budget(analysis)
        outline_mode = analysis.get("outline_mode")
        chapter_titles = {self._clean_compare_text(ch.get("title", "")) for ch in chapters if isinstance(ch, dict)}
        cleaned_chapters = []
        global_seen = {}

        for chapter in chapters:
            if not isinstance(chapter, dict):
                continue
            title = (chapter.get("title") or "").strip()
            if not title or self._is_noise_line(title):
                continue

            cleaned_kps = []
            local_seen = set()
            for kp in chapter.get("knowledge_points") or []:
                kp_data = kp if isinstance(kp, dict) else {"title": str(kp), "description": ""}
                kp_title = str(kp_data.get("title", "") or "").strip() if outline_mode == "java_toc" else self._summarize_to_kp_title(kp_data.get("title", ""))
                compare = self._clean_compare_text(kp_title)
                if (
                    not kp_title
                    or compare in local_seen
                    or compare in chapter_titles
                    or self._is_pseudo_knowledge_point(kp_title)
                ):
                    continue
                local_seen.add(compare)
                global_seen.setdefault(compare, title)
                description = self._build_kp_description(kp_title, kp_data.get("description"), title)
                cleaned_kps.append({
                    **kp_data,
                    "title": kp_title,
                    "description": description,
                    "category": kp_data.get("category") or "核心知识点",
                })

            if not cleaned_kps:
                cleaned_kps = self._derive_kps_from_chapter(chapter)
                for kp in cleaned_kps:
                    global_seen.setdefault(self._clean_compare_text(kp["title"]), title)

            cleaned_chapters.append({
                **chapter,
                "title": title,
                "description": chapter.get("description") or "",
                "knowledge_points": cleaned_kps[:self._chapter_kp_limit(chapter, budget)],
            })

        analysis["chapters"] = cleaned_chapters
        analysis["objectives"] = self._normalize_text_list(analysis.get("objectives"))[:12]
        analysis["prerequisites"] = self._normalize_text_list(analysis.get("prerequisites"))[:10]
        analysis["references"] = self._normalize_text_list(analysis.get("references"))[:12]
        return analysis

    def _clean_compare_text(self, text):
        text = re.sub(r"^\d+(?:\.\d+)*\s*", "", str(text or ""))
        text = re.sub(r"[\s:：，。、；、\-\[\]（）\(\)]", "", text)
        return text.lower()

    def _is_pseudo_knowledge_point(self, text):
        value = str(text or "").strip()
        if not value:
            return True
        if re.search(r"\[[JMD]\]|^\[\d+\]", value, re.I):
            return True
        pseudo_words = ["备注", "注意项", "参考资料", "附录", "附件", "其他说明", "补充说明", "注", "附"]
        if any(word == value or value.startswith(word) for word in pseudo_words):
            return True
        return False

    def _is_garbled_pdf_line(self, text):
        value = str(text or "").strip()
        if not value:
            return True
        meaningful_chars = re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", value)
        if not meaningful_chars:
            return True
        non_ascii = [ch for ch in value if ord(ch) > 127]
        cjk = re.findall(r"[\u4e00-\u9fff]", value)
        if non_ascii and not cjk:
            return True
        if len(non_ascii) >= 2 and not cjk and len(re.findall(r"[A-Za-z0-9]", value)) < len(value) * 0.45:
            return True
        return False

    def _is_heading_noise_title(self, title):
        value = str(title or "").strip()
        if not value:
            return True
        if re.match(r"^(实例\s*\d*|注意|注意：|Gif 图演示|代码：?|.*文件代码：?)$", value, re.I):
            return True
        if re.match(r"^(语法|示例|代码)$", value, re.I):
            return True
        if re.search(r"\.java\s*文件代码", value, re.I):
            return True
        if len(value) > 1 and re.search(r"[{};=]", value):
            return True
        if re.search(r"(http|www\.|\.com|点击|如下图|图片|下载地址)", value, re.I):
            return True
        if value.lower() in {"youkeda", "path设置"}:
            return True
        return False

    def _derive_kps_from_chapter(self, chapter):
        title = chapter.get("title") or ""
        description = chapter.get("description") or ""
        candidates = []
        title_clean = re.sub(r"^\d+(?:\.\d+)*\s*", "", title).strip()
        if title_clean and not self._is_pseudo_knowledge_point(title_clean):
            candidates.append(title_clean)
        candidates.extend(self._extract_kp_candidates_from_line(description))
        result = []
        seen = set()
        for candidate in candidates:
            kp_title = self._summarize_to_kp_title(candidate)
            compare = self._clean_compare_text(kp_title)
            if kp_title and compare not in seen and not self._is_pseudo_knowledge_point(kp_title):
                seen.add(compare)
                result.append({"title": kp_title, "description": description[:180], "category": "核心知识点"})
        return result

    def _build_kp_description(self, title, description="", chapter_title=""):
        desc = re.sub(r"\s+", " ", str(description or "")).strip()
        if desc:
            return desc[:240]
        title = re.sub(r"\s+", " ", str(title or "")).strip()
        chapter_title = re.sub(r"\s+", " ", str(chapter_title or "")).strip()
        if chapter_title and title and self._clean_compare_text(chapter_title) != self._clean_compare_text(title):
            return f"{chapter_title}中的核心知识点：{title}"
        if title:
            return f"课程大纲中提取的核心知识点：{title}"
        return ""

    def _rule_based_parse(self, text, course):
        # 检测 PDF/DOCX 提取的目录结构标记
        toc_text, main_text = None, None
        toc_marker = "【文档目录结构】"
        content_marker = "【正文内容】"
        if toc_marker in (text or "") and content_marker in (text or ""):
            toc_idx = text.index(toc_marker) + len(toc_marker)
            content_idx = text.index(content_marker)
            toc_text = text[toc_idx:content_idx].strip()
            main_text = text[content_idx + len(content_marker):].strip()
            logger.info(f"检测到文档目录结构标记，目录行数：{len(toc_text.splitlines())}，正文行数：{len((main_text or '').splitlines())}")

        raw_lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
        heading_result = self._parse_marked_heading_structure(raw_lines, course, text)
        if heading_result:
            return heading_result

        lines = [
            line.strip(" \t\r\n#-*·、")
            for line in raw_lines
            if not self._is_garbled_pdf_line(line)
        ]
        toc_result = self._parse_java_toc_outline(lines, course, text)
        if toc_result:
            return toc_result

        objectives, prereqs, refs = [], [], []
        chapters = []
        current = None
        objective_words = ["教学目标", "课程目标", "学习目标"]
        prereq_words = ["先修", "前置", "先修课程"]
        ref_words = ["参考文献", "参考", "引用"]

        def is_chapter_heading(value):
            value = value.strip()
            lead_words = [
                "\u77e5\u8bc6\u70b9", "\u91cd\u70b9", "\u96be\u70b9", "\u6838\u5fc3\u5185\u5bb9",
                "\u6559\u5b66\u5185\u5bb9", "\u5b66\u4e60\u5185\u5bb9", "\u638c\u63e1", "\u7406\u89e3",
                "\u719f\u6089", "\u4e86\u89e3",
            ]
            if re.match(r"^(" + "|".join(re.escape(w) for w in lead_words) + r")\s*[:\uff1a]?", value):
                return False
            if re.match(r"^(\d+(?:\.\d+)*\s+.{2,60}|chapter\s*\d+|unit\s*\d+|module\s*\d+)", value, re.I):
                return True
            return 2 <= len(value) <= 36 and not re.search(r"[.!?;]", value) and self._knowledge_signal_score(value) >= 1

        # 如果存在目录结构，先基于目录构建章节骨架
        toc_chapters = []
        if toc_text:
            toc_chapters = self._parse_toc_structure(toc_text)

        # 使用目录骨架或正文逐行解析
        if toc_chapters:
            chapters = toc_chapters
            # 用正文内容丰富章节描述和知识点
            if main_text:
                main_lines = [line.strip(" \t\r\n#-*·、") for line in main_text.splitlines() if line.strip()]
                self._enrich_chapters_from_main_text(chapters, main_lines, objectives, objective_words, prereqs, prereq_words, refs, ref_words)
        else:
            for line in lines:
                if self._is_noise_line(line):
                    continue
                if any(word in line for word in objective_words):
                    objectives.extend(self._split_items(line))
                    continue
                if any(word in line for word in prereq_words):
                    prereqs.extend(self._split_items(line))
                    continue
                if any(word in line for word in ref_words):
                    refs.append(line)
                    continue

                if is_chapter_heading(line):
                    if current:
                        self._finalize_rule_chapter(current)
                    current = {"title": line, "description": "", "knowledge_points": [], "teaching_hours": 0, "chapter_type": "theory"}
                    title_kp = self._knowledge_point_from_heading(line)
                    if title_kp:
                        current["knowledge_points"].append(title_kp)
                    chapters.append(current)
                    continue

                if current:
                    items = self._extract_kp_candidates_from_line(line)
                    existing = {kp.get("title") for kp in current["knowledge_points"]}
                    for item in items:
                        title = self._summarize_to_kp_title(item)
                        if title and title not in existing:
                            current["knowledge_points"].append({"title": title, "description": item if item != title else "", "category": "核心知识点"})
                            existing.add(title)
                    if items:
                        continue
                    if len(line) > 18 and not current["description"]:
                        current["description"] = line
                        continue

        if current:
            self._finalize_rule_chapter(current)

        if not chapters:
            chunks = self._extract_candidate_terms(text)
            if not chunks:
                chunks = self._split_items("；".join(lines))
            fallback_budget = self._build_analysis_budget(text, {
                "chapters": [{"title": "综合概述", "knowledge_points": [{"title": item} for item in chunks]}]
            })
            chapters.append({
                "title": "课程大纲",
                "description": (text or "")[:300],
                "knowledge_points": [{"title": item, "description": "", "category": "核心知识点"} for item in chunks[:fallback_budget["total_kp_limit"]]],
                "teaching_hours": 0,
                "chapter_type": "theory",
            })

        result = {
            "course": {"title": getattr(course, "title", "课程"), "description": getattr(course, "description", "")},
            "objectives": [o for o in objectives if len(o) > 2][:12],
            "prerequisites": [p for p in prereqs if len(p) > 1][:10],
            "references": refs[:12],
            "chapters": chapters,
            "raw_text": text or "",
        }
        result["_analysis_budget"] = self._build_analysis_budget(text, result)
        return result

    def _parse_marked_heading_structure(self, raw_lines, course, raw_text):
        """Build a document outline from heading markers injected by PDF/DOCX extraction."""
        heading_pattern = re.compile(r"^(#{1,6})\s+(.+)$")
        heading_count = sum(1 for line in raw_lines if heading_pattern.match(line))
        if heading_count < 3:
            return None

        chapters = []
        current_chapter = None
        current_kp = None
        description_lines = []

        def flush_description():
            nonlocal description_lines, current_chapter, current_kp
            if not description_lines:
                return
            description = " ".join(description_lines).strip()
            description_lines = []
            if not description:
                return
            if current_kp and not current_kp.get("description"):
                current_kp["description"] = description[:240]
            elif current_chapter and not current_chapter.get("description"):
                current_chapter["description"] = description[:300]

        def add_kp(chapter, title, description=""):
            title = self._summarize_to_kp_title(title)
            if not title or self._is_heading_noise_title(title):
                return None
            existing = {
                self._clean_compare_text(kp.get("title", "") if isinstance(kp, dict) else str(kp))
                for kp in chapter.get("knowledge_points", [])
            }
            compare = self._clean_compare_text(title)
            if not compare or compare in existing:
                return None
            kp = {"title": title, "description": description[:240], "category": "核心知识点"}
            chapter.setdefault("knowledge_points", []).append(kp)
            return kp

        meaningful_heading_levels = []
        for line in raw_lines:
            match = heading_pattern.match(line)
            if not match:
                continue
            title = re.sub(r"^\d+(?:\.\d+)*[。．]?\s*", "", match.group(2).strip()).strip()
            if (
                title
                and not self._is_noise_line(title)
                and not self._is_garbled_pdf_line(title)
                and not self._is_heading_noise_title(title)
            ):
                meaningful_heading_levels.append(len(match.group(1)))
        chapter_level = None
        if meaningful_heading_levels:
            level_counts = defaultdict(int)
            for level in meaningful_heading_levels:
                level_counts[level] += 1
            levels = sorted(level_counts)
            chapter_level = levels[0]
            if len(levels) > 1 and level_counts[chapter_level] < 3:
                chapter_level = levels[1]

        for line in raw_lines:
            match = heading_pattern.match(line)
            if not match:
                if current_chapter:
                    clean_line = self._cleanup_extracted_text(line)
                    if (
                        clean_line
                        and not self._is_noise_line(clean_line)
                        and not self._is_garbled_pdf_line(clean_line)
                        and len(clean_line) > 8
                    ):
                        description_lines.append(clean_line)
                        if len(description_lines) >= 3:
                            flush_description()
                continue

            flush_description()
            level = len(match.group(1))
            title = match.group(2).strip()
            title = re.sub(r"^\d+(?:\.\d+)*[。．]?\s*", "", title).strip()
            if (
                not title
                or self._is_noise_line(title)
                or self._is_garbled_pdf_line(title)
                or self._is_heading_noise_title(title)
            ):
                continue

            if level <= chapter_level:
                current_chapter = {
                    "title": title,
                    "description": "",
                    "knowledge_points": [],
                    "teaching_hours": 0,
                    "chapter_type": "theory",
                }
                chapters.append(current_chapter)
                current_kp = add_kp(current_chapter, title)
            elif current_chapter:
                current_kp = add_kp(current_chapter, title)
            else:
                current_chapter = {
                    "title": title,
                    "description": "",
                    "knowledge_points": [],
                    "teaching_hours": 0,
                    "chapter_type": "theory",
                }
                chapters.append(current_chapter)
                current_kp = add_kp(current_chapter, title)

        flush_description()
        valid_chapters = []
        for chapter in chapters:
            self._finalize_rule_chapter(chapter)
            if chapter.get("knowledge_points"):
                valid_chapters.append(chapter)

        if len(valid_chapters) < 2 or self._count_knowledge_points(valid_chapters) < 4:
            return None

        result = {
            "course": {"title": getattr(course, "title", "课程"), "description": getattr(course, "description", "")},
            "objectives": [],
            "prerequisites": [],
            "references": [],
            "chapters": valid_chapters,
            "raw_text": raw_text or "",
        }
        result["_analysis_budget"] = self._build_analysis_budget(raw_text, result)
        return result

    def _parse_toc_structure(self, toc_text):
        """解析目录结构文本，构建章节骨架。目录行格式：# 标题、## 标题 等"""
        chapters = []
        current_chapter = None
        current_sub_kps = []

        for line in toc_text.splitlines():
            line = line.strip()
            if not line:
                continue
            # 匹配 Markdown 标题层级标记：# 标题、## 标题
            m = re.match(r'^(#{1,6})\s+(.+)$', line)
            if m:
                level = len(m.group(1))
                title = m.group(2).strip()
                if not title or self._is_noise_line(title):
                    continue
                if level == 1:
                    # 一级标题作为章节
                    if current_chapter:
                        # 将子级知识点合并到当前章节
                        if current_sub_kps:
                            existing = {kp.get("title") for kp in current_chapter["knowledge_points"]}
                            for kp_title in current_sub_kps:
                                if kp_title not in existing:
                                    current_chapter["knowledge_points"].append({"title": kp_title, "description": "", "category": "核心知识点"})
                                    existing.add(kp_title)
                        self._finalize_rule_chapter(current_chapter)
                    current_chapter = {"title": title, "description": "", "knowledge_points": [], "teaching_hours": 0, "chapter_type": "theory"}
                    title_kp = self._knowledge_point_from_heading(title)
                    if title_kp:
                        current_chapter["knowledge_points"].append(title_kp)
                    chapters.append(current_chapter)
                    current_sub_kps = []
                elif level >= 2 and current_chapter:
                    # 二级及以下标题作为知识点
                    kp_title = re.sub(r'^\d+(?:\.\d+)*[。．]?\s*', '', title).strip()
                    if kp_title and self._is_meaningful_title(kp_title):
                        current_sub_kps.append(kp_title)
            else:
                # 无标题标记的行，尝试作为当前章节的子知识点
                title = re.sub(r'^\d+(?:\.\d+)*[。．]?\s*', '', line).strip()
                if title and self._is_meaningful_title(title):
                    if current_chapter:
                        current_sub_kps.append(title)

        # 处理最后一个章节
        if current_chapter and current_sub_kps:
            existing = {kp.get("title") for kp in current_chapter["knowledge_points"]}
            for kp_title in current_sub_kps:
                if kp_title not in existing:
                    current_chapter["knowledge_points"].append({"title": kp_title, "description": "", "category": "核心知识点"})
                    existing.add(kp_title)
        if current_chapter:
            self._finalize_rule_chapter(current_chapter)

        return chapters

    def _enrich_chapters_from_main_text(self, chapters, main_lines, objectives, objective_words, prereqs, prereq_words, refs, ref_words):
        """用正文内容丰富目录骨架章节的描述和知识点"""
        # 构建章节标题到章节的映射
        chapter_map = {}
        for ch in chapters:
            chapter_map[ch["title"].strip().lower()] = ch
            # 也支持部分匹配
            short_title = re.sub(r'^\d+(?:\.\d+)*[。．]?\s*', '', ch["title"]).strip().lower()
            if short_title != ch["title"].strip().lower():
                chapter_map[short_title] = ch

        current_chapter = None
        for line in main_lines:
            if self._is_noise_line(line):
                continue
            # 检查是否匹配到某个章节标题
            matched_chapter = None
            line_lower = line.strip().lower()
            for title_key, ch in chapter_map.items():
                if title_key in line_lower or line_lower in title_key:
                    matched_chapter = ch
                    break

            if matched_chapter:
                current_chapter = matched_chapter
                # 如果章节还没有描述，用这行作为描述
                if not current_chapter["description"] and len(line) > 10:
                    current_chapter["description"] = line
                continue

            # 检查目标、先修、参考文献
            if any(word in line for word in objective_words):
                objectives.extend(self._split_items(line))
                continue
            if any(word in line for word in prereq_words):
                prereqs.extend(self._split_items(line))
                continue
            if any(word in line for word in ref_words):
                refs.append(line)
                continue

            # 为当前章节补充知识点和描述
            if current_chapter:
                items = self._extract_kp_candidates_from_line(line)
                existing = {kp.get("title") for kp in current_chapter["knowledge_points"]}
                for item in items:
                    title = self._summarize_to_kp_title(item)
                    if title and title not in existing:
                        current_chapter["knowledge_points"].append({"title": title, "description": item if item != title else "", "category": "核心知识点"})
                        existing.add(title)
                if items:
                    continue
                if len(line) > 18 and not current_chapter["description"]:
                    current_chapter["description"] = line
                    continue

    def _parse_java_toc_outline(self, lines, course, raw_text):
        def u(value):
            return value.encode("ascii").decode("unicode_escape") if "\\u" in value else value

        main_headings = {
            "Java " + u("\u5165\u95e8"),
            "Java " + u("\u9762\u5411\u5bf9\u8c61"),
            "Java " + u("\u96c6\u5408"),
            "Java " + u("\u591a\u7ebf\u7a0b"),
            "Java " + u("\u9ad8\u7ea7\u8bed\u6cd5"),
            "Java " + u("\u65b0\u7279\u6027"),
        }
        chapters = []
        current = None

        for raw_line in lines:
            line = re.sub(r"\s+", " ", str(raw_line or "")).strip(" -\t\r\n")
            line = re.sub(r"^[\u25cb\u25cf\u25e6\u25aa\u25ab]\s*", "", line).strip()
            if not line:
                continue

            if line in main_headings:
                current = {
                    "title": line,
                    "description": line + u("\u76f8\u5173\u77e5\u8bc6\u7ebf\u8def"),
                    "knowledge_points": [],
                    "teaching_hours": 0,
                    "chapter_type": "theory",
                }
                chapters.append(current)
                continue

            if current and re.match(r"^Java\s+", line, re.I):
                title = self._clean_java_toc_title(line)
                if title and title != current["title"]:
                    current["knowledge_points"].append({
                        "title": title,
                        "description": current["title"] + u("\u7ebf\u8def\u4e0b\u7684\u4e8c\u7ea7\u77e5\u8bc6\u70b9\uff1a") + title,
                        "category": u("\u4e8c\u7ea7\u77e5\u8bc6\u70b9"),
                    })

        valid_chapters = [ch for ch in chapters if ch.get("knowledge_points")]
        if len(valid_chapters) < 4:
            return None

        result = {
            "course": {"title": getattr(course, "title", "Java " + u("\u6559\u8f85\u8d44\u6599")), "description": getattr(course, "description", "")},
            "objectives": [],
            "prerequisites": [],
            "references": [],
            "chapters": valid_chapters,
            "raw_text": raw_text or "",
            "outline_mode": "java_toc",
        }
        result["_analysis_budget"] = self._build_analysis_budget(raw_text, result)
        return result

    def _clean_java_toc_title(self, line):
        title = re.sub(r"\s+", " ", str(line or "")).strip()
        title = title.replace("?", "(").replace("?", ")")
        title = re.sub(r"^[\d.?\-\s]+", "", title).strip()
        if not title or len(title) > 80:
            return ""
        return title

    def _is_noise_line(self, line):
        text = (line or "").strip()
        if not text or len(text) <= 1:
            return True
        date_chars = "\u5e74\u6708\u65e5".encode().decode("unicode_escape")
        if re.fullmatch(r"[_\-\s\d/:：()（）．。?" + re.escape(date_chars) + r"]+", text):
            return True
        noise_words = ["\u59d3\u540d", "\u5b66\u53f7", "\u4e13\u4e1a", "\u6307\u5bfc\u6559\u5e08", "\u8bc4\u9605\u4eba", "\u590d\u6838\u4eba", "\u5f97\u5206", "\u76ee\u5f55"]
        noise_words = [w.encode().decode("unicode_escape") for w in noise_words]
        if any(word in text for word in noise_words):
            return True
        page_char = "\u9875".encode().decode("unicode_escape")
        prefix_char = "\\u7b2c".encode().decode("unicode_escape")
        if re.search(prefix_char + r"\s*\d+\s*" + page_char, text):
            return True
        end_words = ["\u6458\u8981", "\u5173\u952e\u8bcd", "\u53c2\u8003\u6587\u732e", "\u81f4\u8c22"]
        end_words = [w.encode().decode("unicode_escape") for w in end_words]
        if any(text.endswith(word) for word in end_words) and len(text) <= 12:
            return True
        return False

    def _knowledge_point_from_heading(self, line):
        text = (line or "").strip()
        if not text:
            return None
        cleaned = re.sub(r"^\d+(?:\.\d+)*[。．]?\s*", "", text)
        cn_nums = "\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e".encode().decode("unicode_escape")
        cleaned = re.sub(r"^[" + cn_nums + r"]+[。．]?\s*", "", cleaned)
        parts = re.split(r"[：:]", cleaned, maxsplit=1)
        title = parts[0].strip(" -：")
        description = parts[1].strip() if len(parts) > 1 else cleaned
        if len(title) > 48:
            title = title[:48].strip()
        banned = ["\u9879\u76ee\u65b9\u6848", "\u7814\u7a76\u80cc\u666f", "\u5b9e\u65bd\u8ba1\u5212", "\u7ecf\u8d39\u9884\u7b97", "\u56e2\u961f\u6210\u5458", "\u7533\u62a5\u57fa\u7840"]
        banned = [w.encode().decode("unicode_escape") for w in banned]
        if len(title) < 2 or self._is_noise_line(title) or self._is_pseudo_knowledge_point(title) or title in banned:
            return None
        return {
            "title": title,
            "description": description[:240] if description != title else "",
            "category": "核心知识点",
        }

    def _extract_kp_candidates_from_line(self, line):
        text = (line or "").strip()
        if not text or self._is_noise_line(text):
            return []
        lead_words = [
            "\u77e5\u8bc6\u70b9", "\u91cd\u70b9", "\u96be\u70b9", "\u6838\u5fc3\u5185\u5bb9",
            "\u6559\u5b66\u5185\u5bb9", "\u5b66\u4e60\u5185\u5bb9", "\u638c\u63e1", "\u7406\u89e3",
            "\u719f\u6089", "\u4e86\u89e3", "\u80fd\u591f", "\u4f1a",
        ]
        has_explicit_marker = bool(re.match(r"^(" + "|".join(re.escape(w) for w in lead_words) + r")\s*[:\uff1a]\s*", text))
        text = re.sub(r"^(" + "|".join(re.escape(w) for w in lead_words) + r")\s*[:\uff1a]\s*", "", text)
        items = self._split_items(text)
        scored = []
        for item in items or [text]:
            item = item.strip(" -:\uff1a()\uff08\uff09[]\u3010\u3011")
            score = self._knowledge_signal_score(item)
            if score >= 2 or (has_explicit_marker and 2 <= len(item) <= 56 and not self._is_noise_line(item)):
                scored.append(item)
        return scored[:6]

    def _knowledge_signal_score(self, text):
        if not text:
            return 0
        text = str(text).strip()
        if self._is_noise_line(text):
            return 0
        score = 0
        if 2 <= len(text) <= 56:
            score += 1
        if re.match(r"^[A-Za-z][A-Za-z0-9_+\-#./ ]{1,40}$", text):
            score += 1
        concept_words = [
            "\u6982\u5ff5", "\u539f\u7406", "\u65b9\u6cd5", "\u6280\u672f", "\u7406\u8bba",
            "\u6a21\u578b", "\u7b97\u6cd5", "\u6846\u67b6", "\u673a\u5236", "\u7b56\u7565",
            "\u4f53\u7cfb", "\u7ed3\u6784", "\u6a21\u5f0f", "\u89c4\u8303", "\u6807\u51c6",
            "\u8981\u7d20", "\u7279\u5f81", "\u77e5\u8bc6\u70b9", "Java", "Python",
            "\u7c7b", "\u5bf9\u8c61", "\u63a5\u53e3", "\u7ee7\u627f", "\u51fd\u6570",
            "\u53d8\u91cf", "\u6570\u7ec4", "\u96c6\u5408", "\u7ebf\u7a0b", "\u5f02\u5e38",
            "\u6570\u636e\u5e93", "\u7f51\u7edc",
        ]
        action_words = [
            "\u8bbe\u8ba1", "\u5b9e\u73b0", "\u5206\u6790", "\u4f18\u5316", "\u8bc4\u4f30",
            "\u90e8\u7f72", "\u6d4b\u8bd5", "\u8c03\u8bd5", "\u914d\u7f6e", "\u7ba1\u7406",
            "\u5f00\u53d1", "\u5e94\u7528", "\u4f7f\u7528", "\u6784\u5efa",
        ]
        learning_words = ["\u638c\u63e1", "\u7406\u89e3", "\u719f\u6089", "\u4e86\u89e3", "\u80fd\u591f", "\u5b66\u4f1a", "\u8ba4\u8bc6"]
        if any(word in text for word in concept_words):
            score += 2
        if any(word in text for word in action_words):
            score += 1
        if any(word in text for word in learning_words):
            score += 1
        if re.search(r"[\u3002\uff01?\uff1b]$", text) and len(text) > 45:
            score -= 1
        if len(text) > 90:
            score -= 2
        return score

    def _summarize_to_kp_title(self, text):
        item = re.sub(r"\s+", " ", str(text or "")).strip(" -:\uff1a;\uff1b,\uff0c")
        if not item or self._knowledge_signal_score(item) < 1:
            return ""
        item = re.sub(r"^\d+(?:\.\d+)*\s*", "", item)
        topic_words = [
            "\u77e5\u8bc6\u70b9", "\u91cd\u70b9", "\u96be\u70b9", "\u6838\u5fc3\u5185\u5bb9",
            "\u6559\u5b66\u5185\u5bb9", "\u5b66\u4e60\u5185\u5bb9",
        ]
        learning_words = ["\u638c\u63e1", "\u7406\u89e3", "\u719f\u6089", "\u4e86\u89e3", "\u80fd\u591f", "\u4f1a"]
        prefixes = ["\u7b2c", "\u6ce8", "\u9644", "\u5907\u6ce8", "\u6ce8\u610f", "\u8bf4\u660e", "\u53c2\u8003", "\u5176\u4ed6"]
        item = re.sub(r"^(" + "|".join(re.escape(w) for w in topic_words) + r")\s*[:\uff1a]\s*", "", item)
        item = re.sub(r"^(" + "|".join(re.escape(w) for w in learning_words) + r")\s*", "", item)
        for prefix in prefixes:
            if item.startswith(prefix):
                item = item[len(prefix):].strip()
        if len(item) <= 56:
            return item
        parts = re.split(r"[\uff0c\u3002 \u3001\uff1b;:\uff1a]", item)
        for part in parts:
            part = part.strip()
            if 2 <= len(part) <= 56 and self._knowledge_signal_score(part) >= 1:
                return part
        return item[:56].strip()

    def _finalize_rule_chapter(self, chapter):
        seen = set()
        cleaned = []
        for kp in chapter.get("knowledge_points", []):
            title = self._summarize_to_kp_title(kp.get("title", "") if isinstance(kp, dict) else str(kp))
            if not title or title in seen:
                continue
            seen.add(title)
            cleaned.append({
                "title": title,
                "description": (kp.get("description", "") if isinstance(kp, dict) else "")[:180],
                "category": kp.get("category", "核心知识点") if isinstance(kp, dict) else "核心知识点",
            })
        if not cleaned and chapter.get("description"):
            for candidate in self._extract_kp_candidates_from_line(chapter["description"]):
                title = self._summarize_to_kp_title(candidate)
                if title and title not in seen:
                    seen.add(title)
                    cleaned.append({"title": title, "description": candidate, "category": "核心知识点"})
        chapter["knowledge_points"] = cleaned

    def _normalize_kp_list(self, key_points):
        result = []
        for kp in key_points or []:
            if isinstance(kp, str):
                result.append({"title": kp, "description": "", "category": "核心知识点"})
            elif isinstance(kp, dict):
                result.append({
                    "title": str(kp.get("title") or kp.get("name") or ""),
                    "description": str(kp.get("description") or kp.get("definition") or ""),
                    "category": str(kp.get("category") or "核心知识点"),
                    "related_concepts": kp.get("related_concepts") or kp.get("related") or [],
                })
        return [kp for kp in result if kp["title"]]

    def _normalize_text_list(self, values):
        if values is None:
            return []
        if isinstance(values, str):
            values = self._split_items(values) or [values]
        result = []
        for item in values:
            if isinstance(item, dict):
                text = (
                    item.get("title")
                    or item.get("name")
                    or item.get("objective")
                    or item.get("description")
                    or item.get("content")
                    or item.get("text")
                    or ""
                )
            else:
                text = item
            text = str(text or "").strip()
            if text and text not in result:
                result.append(text)
        return result

    # Long text chunking configuration.
    LITE_CONTEXT_LIMIT = 10000  # 单块上下文上限：较小分块降低 Spark prefill 延迟，配合 bulkhead=2 总耗时更优
    CHUNK_OVERLAP = 200
    MIN_CHUNK_SIZE = 1500
    MAX_CONCURRENT_WORKERS = 4
    MIN_CONCURRENT_WORKERS = 1

    def _llm_based_parse(self, text, course, user_id=None, user_role=None):
        """Use Spark LLM for structured extraction, with long text chunking."""
        from src.services.spark_service import spark_service

        text = text or ""
        if len(text) <= self.LITE_CONTEXT_LIMIT:
            return self._llm_parse_single(text, course, user_id=user_id, user_role=user_role)

        logger.info("长文本分块并发解析: 文本长度=%d, 上下文上限=%d", len(text), self.LITE_CONTEXT_LIMIT)
        chunks = self._split_text_for_llm(text)
        logger.info("分为 %d 个片段，开始并发解析", len(chunks))

        max_workers = self._calc_concurrent_workers(len(text), len(chunks))
        logger.info("动态并发线程数: %d", max_workers)

        chunk_results = self._concurrent_parse_chunks(chunks, course, max_workers, user_id=user_id, user_role=user_role)

        if not chunk_results:
            logger.warning("所有片段解析失败，回退到截断解析")
            return self._llm_parse_single(text[:self.LITE_CONTEXT_LIMIT], course, user_id=user_id, user_role=user_role)

        merged = self._merge_chunk_results(chunk_results, course)

        merged = self._cross_validate_results(merged)

        logger.info("分段解析完成: %d 个片段成功，合并后 %d 章节 %d 知识点",
                     len(chunk_results),
                     len(merged.get("chapters", [])),
                     sum(len(ch.get("knowledge_points", [])) for ch in merged.get("chapters", [])))
        return merged

    def _split_text_for_llm(self, text):
        """Split long text by chapter-like boundaries before LLM parsing."""
        text = text or ""
        limit = self.LITE_CONTEXT_LIMIT
        overlap = self.CHUNK_OVERLAP
        min_chunk = self.MIN_CHUNK_SIZE

        if len(text) <= limit:
            return [text]

        chunks = []
        chapter_pattern = re.compile(
            r'^(第\s*[一二三四五六七八九十百\d]+\s*[章节篇单元]|'
            r'\d+(:\.\d+)\s+.{2,30}|'
            r'chapter\s*\d+|unit\s*\d+|module\s*\d+)',
            re.MULTILINE | re.IGNORECASE
        )

        splits = []
        last_end = 0
        for match in chapter_pattern.finditer(text):
            if match.start() > last_end:
                splits.append((last_end, match.start()))
            last_end = match.start()
        if last_end < len(text):
            splits.append((last_end, len(text)))

        for start, end in splits:
            segment = text[start:end]
            if len(segment) <= limit:
                chunks.append(segment)
            else:
                sub_chunks = self._split_by_paragraphs(segment, limit, overlap, min_chunk)
                chunks.extend(sub_chunks)

        if not chunks:
            chunks = self._split_by_paragraphs(text, limit, overlap, min_chunk)

        final_chunks = []
        for i, chunk in enumerate(chunks):
            if i > 0 and overlap > 0:
                prev = chunks[i - 1]
                prefix = prev[-overlap:] if len(prev) > overlap else prev
                chunk = prefix + "\n" + chunk
            if len(chunk) > limit:
                chunk = chunk[:limit]
            final_chunks.append(chunk)

        return final_chunks

    def _split_by_paragraphs(self, text, limit, overlap, min_chunk):
        """Split text by paragraph and sentence boundaries."""
        paragraphs = re.split(r'\n\s*\n|\n(=[\d]+[.])', text)
        chunks = []
        current = ""

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            if len(current) + len(para) + 1 <= limit:
                current = current + "\n" + para if current else para
                continue

            if current:
                chunks.append(current)
            if len(para) > limit:
                sentences = re.split(r'(<=[\n])', para)
                sub = ""
                for sent in sentences:
                    if len(sub) + len(sent) <= limit:
                        sub += sent
                    else:
                        if sub:
                            chunks.append(sub)
                        sub = sent[:limit]
                current = sub or ""
            else:
                current = para

        if current and len(current) >= min_chunk:
            chunks.append(current)
        elif current and chunks:
            chunks[-1] = chunks[-1] + "\n" + current
        elif current:
            chunks.append(current)

        return chunks

    def _calc_concurrent_workers(self, text_length, chunk_count):
        """Calculate concurrent LLM workers based on text length and chunk count."""
        import os
        cpu_count = os.cpu_count() or 2

        if text_length > 50000:
            base_workers = min(4, cpu_count)
        elif text_length > 20000:
            base_workers = min(3, cpu_count)
        elif text_length > 8000:
            base_workers = 2
        else:
            base_workers = 1

        workers = min(base_workers, chunk_count, self.MAX_CONCURRENT_WORKERS)
        workers = max(workers, self.MIN_CONCURRENT_WORKERS)
        return workers

    def _concurrent_parse_chunks(self, chunks, course, max_workers, user_id=None, user_role=None):
        """Parse multiple chunks concurrently."""
        results = []
        lock = threading.Lock()

        def parse_one(index, chunk_text):
            try:
                result = self._llm_parse_single(chunk_text, course, user_id=user_id, user_role=user_role)
                if result:
                    with lock:
                        results.append((index, result))
                    logger.info("片段 %d 解析成功", index)
                else:
                    logger.warning("片段 %d 解析返回空结果", index)
            except Exception as e:
                logger.warning("片段 %d 解析异常: %s", index, e)

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = [executor.submit(parse_one, i, chunk) for i, chunk in enumerate(chunks)]
            timeout_per_chunk = 120
            total_timeout = timeout_per_chunk * max(1, len(chunks))
            for future in as_completed(futures, timeout=total_timeout):
                try:
                    future.result()
                except Exception as e:
                    logger.warning("并发解析异常: %s", e)

        results.sort(key=lambda x: x[0])
        return [r[1] for r in results]

    LLM_MAX_RETRIES = 2
    LLM_RETRY_DELAY = 1.0

    def _llm_parse_single(self, text, course, user_id=None, user_role=None):
        """Parse one text chunk with the LLM, with retry on transient failures."""
        from src.services.spark_service import spark_service
        import time as _time

        prompt = f"""你是课程知识图谱分析专家。请从以下文档内容中提取核心知识点并构建结构化数据。

要求：
1. 提取课程标题、描述、学习目标和先修要求
2. 按文档目录结构提取章节，保留层级关系
3. 为每个知识点提炼核心重点描述（一句话概括关键内容）
4. 识别知识点间的关联关系
5. 严格按以下JSON格式输出，不要包含Markdown代码块标记

JSON格式：
{{
  "title": "课程标题",
  "description": "课程描述",
  "objectives": ["学习目标"],
  "prerequisites": ["先修要求"],
  "chapters": [
    {{
      "title": "章节标题",
      "description": "章节核心内容概述",
      "teaching_hours": 0,
      "chapter_type": "theory",
      "knowledge_points": [
        {{
          "title": "知识点标题",
          "description": "核心重点描述",
          "category": "分类",
          "difficulty": "beginner/intermediate/advanced",
          "related_concepts": ["相关概念"],
          "applies_to": ["应用场景"],
          "assessed_by": ["评估方式"]
        }}
      ]
    }}
  ],
  "cross_chapter_relations": [
    {{"source": "A", "target": "B", "relation": "prerequisite/related/applies_to", "strength": 0.8}}
  ]
}}

文档内容：
{text[:self.LITE_CONTEXT_LIMIT]}"""

        last_error = ""
        for attempt in range(1, self.LLM_MAX_RETRIES + 2):
            try:
                response = spark_service.chat(
                    messages=[{"role": "user", "content": prompt}],
                    user_id=user_id,
                    user_role=user_role,
                    call_type="knowledge_graph_parse",
                )
            except Exception as exc:
                last_error = f"AI请求失败: {exc}"
                logger.warning("LLM解析第%d次尝试请求异常: %s", attempt, exc)
                if attempt <= self.LLM_MAX_RETRIES:
                    _time.sleep(self.LLM_RETRY_DELAY)
                continue

            content = response or ""
            json_match = re.search(r'\{[\s\S]*\}', content)
            if not json_match:
                last_error = "AI返回内容中未找到JSON"
                logger.warning("AI返回内容中未找到JSON (第%d次): %s", attempt, content[:200])
                if attempt <= self.LLM_MAX_RETRIES:
                    _time.sleep(self.LLM_RETRY_DELAY)
                continue
            json_str = json_match.group()
            try:
                data = json.loads(json_str)
            except json.JSONDecodeError:
                data = self._repair_and_parse_json(json_str)
            if not data:
                last_error = "JSON解析失败"
                logger.warning("JSON解析失败 (第%d次)，AI返回内容: %s", attempt, content[:500])
                if attempt <= self.LLM_MAX_RETRIES:
                    _time.sleep(self.LLM_RETRY_DELAY)
                continue
            self._cross_chapter_relations = data.get("cross_chapter_relations", [])
            return self._normalize_structured(json.dumps(data, ensure_ascii=False), course)

        self._last_parse_error = last_error
        logger.error("LLM解析全部重试失败，最后错误: %s", last_error)
        return None

    def _repair_and_parse_json(self, json_str):
        """修复AI返回的常见JSON格式问题"""
        repaired = json_str
        # 去除markdown代码块标记
        repaired = re.sub(r'```(?:json)?\s*', '', repaired)
        repaired = repaired.replace('```', '')
        # 去除注释
        repaired = re.sub(r'//.*?$', '', repaired, flags=re.MULTILINE)
        repaired = re.sub(r'/\*.*?\*/', '', repaired, flags=re.DOTALL)
        # 修复单引号 -> 双引号
        repaired = re.sub(r"'(\w+)'\s*:", r'"\1":', repaired)
        # 修复值中的单引号 -> 双引号
        repaired = re.sub(r":\s*'([^']*)'", r': "\1"', repaired)
        # 去除尾逗号
        repaired = re.sub(r',\s*([}\]])', r'\1', repaired)
        try:
            return json.loads(repaired)
        except json.JSONDecodeError as e:
            logger.warning("JSON修复后仍解析失败: %s, 原始内容: %s", e, json_str[:300])
            return None

    def _merge_chunk_results(self, chunk_results, course):
        """Merge structured results parsed from multiple chunks."""
        if not chunk_results:
            return None
        if len(chunk_results) == 1:
            return chunk_results[0]

        merged = {
            "course": chunk_results[0].get("course", {}),
            "raw_text": "\n".join(r.get("raw_text", "") for r in chunk_results),
        }

        all_objectives = []
        seen_obj = set()
        for result in chunk_results:
            for obj in result.get("objectives", []):
                norm = self._clean_compare_text(obj)
                if norm and norm not in seen_obj and len(obj) > 2:
                    all_objectives.append(obj)
                    seen_obj.add(norm)
        merged["objectives"] = all_objectives[:15]

        all_prereqs = []
        seen_prereq = set()
        for result in chunk_results:
            for prereq in result.get("prerequisites", []):
                norm = self._clean_compare_text(prereq)
                if norm and norm not in seen_prereq and len(prereq) > 1:
                    all_prereqs.append(prereq)
                    seen_prereq.add(norm)
        merged["prerequisites"] = all_prereqs[:12]

        merged["chapters"] = self._merge_chapters_across_chunks(chunk_results)

        all_relations = []
        for result in chunk_results:
            all_relations.extend(result.get("cross_chapter_relations", []))
        seen_relations = set()
        unique_relations = []
        for rel in all_relations:
            key = f"{rel.get('source', '')}-{rel.get('target', '')}-{rel.get('relation', '')}"
            if key not in seen_relations:
                seen_relations.add(key)
                unique_relations.append(rel)
        self._cross_chapter_relations = unique_relations

        return self._normalize_structured(json.dumps(merged, ensure_ascii=False), course)

    def _merge_chapters_across_chunks(self, chunk_results):
        """Merge chapters extracted from different text chunks."""
        all_chapters = []
        for result in chunk_results:
            all_chapters.extend(result.get("chapters", []))

        if not all_chapters:
            return []

        merged = []
        used_indices = set()

        for i, ch in enumerate(all_chapters):
            if i in used_indices:
                continue

            title_i = self._clean_compare_text(ch.get("title", ""))
            combined_kps = list(ch.get("knowledge_points", []))
            combined_desc = ch.get("description", "") or ""

            for j in range(i + 1, len(all_chapters)):
                if j in used_indices:
                    continue
                other = all_chapters[j]
                title_j = self._clean_compare_text(other.get("title", ""))
                is_same_chapter = bool(
                    title_i and title_j and (
                        title_i in title_j
                        or title_j in title_i
                        or self._title_similarity(title_i, title_j) > 0.7
                    )
                )
                if not is_same_chapter:
                    continue

                existing_kp_titles = {
                    self._clean_compare_text(kp.get("title", "") if isinstance(kp, dict) else str(kp))
                    for kp in combined_kps
                }
                for kp in other.get("knowledge_points", []):
                    kp_title = self._clean_compare_text(kp.get("title", "") if isinstance(kp, dict) else str(kp))
                    if kp_title and kp_title not in existing_kp_titles:
                        combined_kps.append(kp)
                        existing_kp_titles.add(kp_title)

                other_desc = other.get("description", "") or ""
                if other_desc and other_desc not in combined_desc:
                    combined_desc = (combined_desc + " " + other_desc).strip()
                used_indices.add(j)

            merged.append({
                **ch,
                "title": ch.get("title", ""),
                "description": combined_desc[:500],
                "knowledge_points": combined_kps,
            })

        return merged

    def _title_similarity(self, title1, title2):
        """Calculate simple character-set title similarity."""
        if not title1 or not title2:
            return 0.0
        set1 = set(title1)
        set2 = set(title2)
        if not set1 or not set2:
            return 0.0
        intersection = set1 & set2
        union = set1 | set2
        return len(intersection) / len(union)

    def _cross_validate_results(self, merged):
        """Cross validate merged chunk results."""
        chapters = merged.get("chapters", [])
        if not chapters:
            return merged

        # 1. 全局知识点去重
        global_seen = set()
        chapter_titles = {self._clean_compare_text(ch.get("title", "")) for ch in chapters}

        for chapter in chapters:
            cleaned_kps = []
            for kp in chapter.get("knowledge_points", []):
                kp_title = kp.get("title", "") if isinstance(kp, dict) else str(kp)
                compare = self._clean_compare_text(kp_title)
                if compare and compare not in global_seen and compare not in chapter_titles:
                    global_seen.add(compare)
                    cleaned_kps.append(kp)
                elif compare in global_seen and isinstance(kp, dict):
                    # 重复知识点：合并描述信息到首次出现的知识点                    for ch2 in chapters:
                        for existing_kp in ch2.get("knowledge_points", []):
                            existing_title = existing_kp.get("title", "") if isinstance(existing_kp, dict) else str(existing_kp)
                            if self._clean_compare_text(existing_title) == compare and isinstance(existing_kp, dict):
                                # 补充描述
                                if kp.get("description") and not existing_kp.get("description"):
                                    existing_kp["description"] = kp["description"]
                                # 补充关联概念
                                for rc in (kp.get("related_concepts") or []):
                                    if rc not in (existing_kp.get("related_concepts") or []):
                                        existing_kp.setdefault("related_concepts", []).append(rc)
                                break
            chapter["knowledge_points"] = cleaned_kps

        # 2. 校验跨章节关系的一致性
        cross_relations = getattr(self, '_cross_chapter_relations', None) or []
        if cross_relations:
            validated = []
            all_kp_titles = global_seen
            for rel in cross_relations:
                source = rel.get("source", "")
                target = rel.get("target", "")
                if (self._clean_compare_text(source) in all_kp_titles and
                    self._clean_compare_text(target) in all_kp_titles):
                    validated.append(rel)
            self._cross_chapter_relations = validated

        return merged

    def _read_docx_text(self, content, file_id):
        binary = self._to_binary(content, file_id)
        parts = []
        outline_parts = []

        try:
            doc = Document(io.BytesIO(binary))
            for p in doc.paragraphs:
                if not p.text or not p.text.strip():
                    continue
                style_name = (p.style.name if p.style else "") or ""
                text = p.text.strip()
                # 标记标题层级，用于目录结构识别
                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.replace("Heading", "").strip())
                    except (ValueError, AttributeError):
                        level = 1
                    marker = "#" * level
                    parts.append(f"{marker} {text}")
                    outline_parts.append({"level": level, "text": text})
                elif style_name.startswith("List"):
                    parts.append(f"• {text}")
                else:
                    parts.append(text)

            for table in doc.tables:
                if not table.rows:
                    continue
                # 提取表头
                header_cells = [cell.text.strip() for cell in table.rows[0].cells if cell.text and cell.text.strip()]
                if header_cells:
                    parts.append(" | ".join(header_cells))
                # 提取表格内容
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
        except Exception as exc:
            logger.warning(f"python-docx extraction failed: {exc}")

        try:
            parts.extend(self._extract_docx_xml_text(binary))
        except Exception as exc:
            logger.warning(f"DOCX XML extraction failed: {exc}")

        deduped = []
        for part in parts:
            text = self._cleanup_extracted_text(part)
            if text and text not in deduped:
                deduped.append(text)
        return "\n".join(deduped)

    def _read_pdf_text(self, content, file_id):
        """解析 PDF 文件内容为纯文本，提取目录结构和文本层级，自动识别标题"""
        binary = self._to_binary(content, file_id)

        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=binary, filetype="pdf")
            try:
                # 提取PDF书签/目录结构
                outline = doc.get_toc()
                outline_texts = []
                if outline:
                    for level, title, page in outline:
                        marker = "#" * level
                        outline_texts.append(f"{marker} {title}")

                # 使用 dict 模式提取带字体信息的文本
                structured_lines = self._extract_pdf_lines_with_font(doc)
                text = self._build_pdf_text_from_lines(structured_lines)

                # 如果有目录结构，将目录信息前置
                if outline_texts:
                    outline_header = "【文档目录结构】\n" + "\n".join(outline_texts) + "\n【正文内容】\n"
                    text = outline_header + text

                if text and text.strip():
                    return text
            finally:
                doc.close()
        except Exception as exc:
            logger.warning(f"PyMuPDF PDF extraction failed: {exc}")

        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(binary))
            # 尝试提取PDF大纲
            outline_texts = []
            try:
                outline = reader.outline
                if outline:
                    def extract_outline_items(items, level=1):
                        for item in items:
                            if isinstance(item, list):
                                extract_outline_items(item, level + 1)
                            elif hasattr(item, 'title'):
                                marker = "#" * level
                                outline_texts.append(f"{marker} {item.title}")
                    extract_outline_items(outline)
            except Exception:
                pass

            text_parts = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(t for t in text_parts if t and t.strip())

            if outline_texts:
                outline_header = "【文档目录结构】\n" + "\n".join(outline_texts) + "\n【正文内容】\n"
                text = outline_header + text

            if text and text.strip():
                return text
        except Exception as exc:
            logger.warning(f"pypdf PDF extraction failed: {exc}")

        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(io.BytesIO(binary)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
                    if page_text.strip():
                        text_parts.append(page_text)
            text = "\n".join(text_parts)
            if text and text.strip():
                return text
        except Exception as exc:
            logger.warning(f"pdfplumber PDF extraction failed: {exc}")

        raise ValueError("未能从PDF中提取到文本，请确认PDF不是扫描图片或已加密")

    def _extract_pdf_lines_with_font(self, doc):
        """从PDF文档中提取带字体大小信息的文本行，并清理页码和重复页眉页脚"""
        from collections import Counter

        all_lines = []  # [(text, font_size, page_num)]
        page_count = doc.page_count

        for page_num, page in enumerate(doc):
            page_dict = page.get_text("dict")
            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:  # 仅处理文本块
                    continue
                for line in block.get("lines", []):
                    line_text = ""
                    total_size = 0.0
                    char_count = 0
                    for span in line.get("spans", []):
                        text = span.get("text", "")
                        if text:
                            line_text += text
                            size = span.get("size", 0)
                            total_size += size * len(text)
                            char_count += len(text)
                    line_text = line_text.strip()
                    if not line_text:
                        continue
                    avg_size = total_size / char_count if char_count > 0 else 0
                    all_lines.append((line_text, avg_size, page_num))

        if not all_lines:
            return []

        # 统计正文字体大小（最常见的）
        size_counter = Counter(round(s, 1) for _, s, _ in all_lines if s > 0)
        body_size = size_counter.most_common(1)[0][0] if size_counter else 10.0

        # 清理页码：纯数字或"第X页"等短行
        cleaned = []
        for text, size, pnum in all_lines:
            stripped = text.strip()
            # 跳过页码（纯数字、纯数字+页/页码等）
            if re.match(r'^\d{1,4}$', stripped):
                continue
            if re.match(r'^第\s*\d+\s*页', stripped):
                continue
            if re.match(r'^[-—\s]*\d+\s*[-—\s]*$', stripped):
                continue
            if self._is_garbled_pdf_line(stripped):
                continue
            cleaned.append((text, size, pnum))

        # 清理重复页眉页脚：在多页中完全相同的短行
        line_page_count = Counter()
        for text, _, _ in cleaned:
            line_page_count[text] += 1
        repeat_threshold = max(3, page_count // 3)
        repeated_lines = {t for t, c in line_page_count.items() if c >= repeat_threshold and len(t) <= 60}

        if repeated_lines:
            cleaned = [(t, s, p) for t, s, p in cleaned if t not in repeated_lines]

        heading_sizes = sorted({round(size, 1) for text, size, _ in cleaned if size >= body_size * 1.2 and len(text) <= 80}, reverse=True)
        major_size = heading_sizes[0] if heading_sizes else body_size * 1.6
        secondary_size = heading_sizes[1] if len(heading_sizes) > 1 else body_size * 1.4

        # 标记标题：字体大小明显大于正文的行
        heading_threshold = body_size * 1.2
        result = []
        for text, size, pnum in cleaned:
            if size >= heading_threshold and len(text) <= 80:
                # 根据字体大小差异确定标题层级
                rounded_size = round(size, 1)
                if rounded_size >= major_size - 0.1:
                    level = 1
                elif rounded_size >= secondary_size - 0.1:
                    level = 2
                else:
                    level = 3
                result.append(("#" * level + " " + text, size, pnum))
            else:
                result.append((text, size, pnum))

        return result

    def _build_pdf_text_from_lines(self, structured_lines):
        """将带字体信息的文本行组装为最终文本"""
        if not structured_lines:
            return ""
        parts = []
        for text, _, _ in structured_lines:
            cleaned = self._cleanup_extracted_text(text)
            if cleaned:
                parts.append(cleaned)
        return "\n".join(parts)

    def _to_binary(self, content, file_id):
        if file_id:
            if hasattr(file_id, "read"):
                pos = None
                try:
                    pos = file_id.tell()
                except Exception:
                    pos = None
                data = file_id.read()
                if pos is not None:
                    try:
                        file_id.seek(pos)
                    except Exception:
                        pass
                return data
            if isinstance(file_id, (bytes, bytearray)):
                return bytes(file_id)
        raw = content or ""
        if isinstance(raw, (bytes, bytearray)):
            return bytes(raw)
        if isinstance(raw, str) and raw.startswith("data:") and "," in raw:
            raw = raw.split(",", 1)[1]
        if isinstance(raw, str):
            try:
                return base64.b64decode(raw, validate=False)
            except Exception:
                return raw.encode("utf-8", errors="ignore")
        return bytes(raw or b"")

    def _extract_docx_xml_text(self, binary):
        with zipfile.ZipFile(io.BytesIO(binary)) as zf:
            text_parts = []
            skip_names = {
                "word/styles.xml",
                "word/settings.xml",
                "word/numbering.xml",
                "word/fontTable.xml",
                "word/webSettings.xml",
            }
            names = [
                name for name in zf.namelist()
                if name.startswith("word/")
                and name.endswith(".xml")
                and name not in skip_names
                and "/_rels/" not in name
            ]
            names.sort(key=lambda name: (name != "word/document.xml", name))
            for name in names:
                text_parts.extend(self._extract_xml_text(zf.read(name)))
            return text_parts

    def _extract_xml_text(self, xml_bytes):
        try:
            root = ET.fromstring(xml_bytes)
        except Exception:
            return []
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        parts = []
        for node in root.findall(".//w:p", ns):
            text = "".join(t.text or "" for t in node.findall(".//w:t", ns))
            text = self._cleanup_extracted_text(text)
            if text:
                parts.append(text)
        return parts

    def _cleanup_extracted_text(self, text):
        cleaned = re.sub(r"\s+", " ", str(text or "")).strip()
        cleaned = cleaned.replace("\u00a0", " ")
        return cleaned

    def _create_source_chunks(self, course_id, normalized):
        chunks = []
        raw_parts = []
        if normalized.get("raw_text"):
            raw_parts.append(("课程大纲", normalized["raw_text"]))
        for chapter in normalized.get("chapters", []):
            title = chapter.get("title", "未知")
            desc = chapter.get("description", "")
            kp_text = "、".join(
                kp.get("title", "") if isinstance(kp, dict) else str(kp)
                for kp in chapter.get("knowledge_points", [])
            )
            raw_parts.append((title, f"{title}\n{desc}\n知识点：{kp_text}"))
        for idx, (title, text) in enumerate(raw_parts, start=1):
            for part_idx, piece in enumerate(self._chunk_text(text), start=1):
                chunk = KnowledgeSourceChunk(
                    course_id=course_id,
                    source_type="syllabus",
                    source_id=f"syllabus:{idx}:{part_idx}",
                    reference_code=f"S{len(chunks) + 1}",
                    title=title,
                    content=piece,
                    location=f"{title} / 片段{part_idx}",
                    metadata_json=json.dumps({"origin": "syllabus_import"}, ensure_ascii=False),
                )
                db.session.add(chunk)
                chunks.append(chunk)
        db.session.flush()
        return chunks

    def _chunk_text(self, text, size=700):
        text = re.sub(r"\s+", " ", text or "").strip()
        if not text:
            return []
        return [text[i:i + size] for i in range(0, len(text), size)]

    def _build_chapter_main_content(self, chapter):
        """构建章节主体内容摘要，融合描述与知识点，去除数字序号前缀"""
        title = chapter.get("title") or "未知章节"
        description = (chapter.get("description") or "").strip()
        kps = [
            kp.get("title") if isinstance(kp, dict) else str(kp)
            for kp in chapter.get("knowledge_points", []) or []
        ]
        kps = [kp for kp in kps if kp and self._is_meaningful_title(kp)]
        # 去除知识点标题中的数字序号前缀
        clean_kps = []
        for kp in kps[:12]:
            cleaned = re.sub(r"^\d+[\.\、\s]+\s*", "", kp).strip()
            if cleaned:
                clean_kps.append(cleaned)
        if clean_kps and description:
            kp_text = "、".join(clean_kps[:10])
            if len(clean_kps) > 10:
                kp_text += f"等共{len(clean_kps)}个知识点"
            return f"{title}：{description[:100]}。核心内容：{kp_text}"
        if clean_kps:
            return f"{title}：" + "、".join(clean_kps[:12])
        if description:
            return f"{title}：{description[:120]}"
        return title

    def _is_meaningful_title(self, title):
        """Return whether a title is meaningful enough for a graph node."""
        if not title or not title.strip():
            return False
        t = title.strip()
        # 过短
        if len(t) < 2:
            return False
        # 纯数字字母格式（如0.0f、.0d、e5等代码字面量）        if re.match(r'^\d+\.\d*[fFdDlL]$', t):
            return False
        # 纯代码标识符（如foo_bar、camelCase等，且过短）
        if re.match(r'^[a-zA-Z_]\w*$', t) and len(t) < 4:
            return False
        # 纯标点或数字
        if re.match(r'^[\d\s\-—,;:;！？。，、：））（】【]+$', t):
            return False
        return True

    def _build_chapter_annotation(self, chapter):
        """基于章节内容生成注释，提炼核心重点信息，去除数字序号前缀避免与标识重复"""
        title = chapter.get("title") or "未知章节"
        description = (chapter.get("description") or "").strip()
        kps = [
            kp.get("title") if isinstance(kp, dict) else str(kp)
            for kp in chapter.get("knowledge_points", []) or []
        ]
        kps = [kp for kp in kps if kp and self._is_meaningful_title(kp)]
        # 去除知识点标题中的数字序号前缀，避免与span标识重复
        clean_kps = []
        for kp in kps[:8]:
            cleaned = re.sub(r"^\d+[\.\、\s]+\s*", "", kp).strip()
            if cleaned:
                clean_kps.append(cleaned)
        if description and clean_kps:
            kp_summary = "、".join(clean_kps[:6])
            if len(clean_kps) > 6:
                kp_summary += f"等共{len(clean_kps)}个知识点"
            return f"{description[:150]}；核心知识点：{kp_summary}"
        if description:
            return description[:260]
        if clean_kps:
            kp_summary = "、".join(clean_kps[:6])
            if len(clean_kps) > 6:
                kp_summary += f"等共{len(clean_kps)}个知识点"
            return f"涵盖：{kp_summary}，建议按目录顺序逐步学习"
        return "待补充详细内容，建议上传包含更丰富章节描述的文档"

    def _split_items(self, text):
        cleaned = str(text or "").strip()
        cleaned = re.sub(r"^[^:\uff1a]{1,12}[:\uff1a]\s*", "", cleaned)
        parts = re.split(r"[;\uff1b,\uff0c\u3001\n\t]|\s{2,}|(?:^|\s)\d+[\u3002.)\u3001]\s*", cleaned)
        results = []
        for part in parts:
            item = part.strip(" -:\uff1a,\uff0c;\uff1b()\uff08\uff09[]\u3010\u3011")
            if 2 <= len(item) <= 80 and item not in results:
                results.append(item)
        return results

    def _extract_candidate_terms(self, text):
        terms = []
        for line in (text or "").splitlines():
            line = line.strip(" \t\r\n#-*·、")
            if not line:
                continue
            if re.search(r"(掌握|理解|了解|熟悉|能够|应用)", line):
                terms.extend(self._split_items(line))
            elif 2 <= len(line) <= 48 and not re.search(r"[。！？]", line):
                terms.append(line)
        deduped = []
        for term in terms:
            if term and term not in deduped:
                deduped.append(term)
        return deduped[:60]

    def _upsert_node(self, graph_id, course_id, node_type, label, description="", category=None, source_chunk_ids=None, properties=None):
        node_type = node_type if node_type in NODE_TYPES else "knowledge_point"
        label = (label or "").strip()[:200]
        node = KnowledgeGraphNode.query.filter_by(course_id=course_id, node_type=node_type, label=label).first()
        is_new = node is None
        if not node:
            node = KnowledgeGraphNode(course_id=course_id, node_type=node_type, label=label, graph_id=graph_id)
            db.session.add(node)
        node.description = description or node.description
        node.category = category or node.category
        node.source_chunk_ids = json.dumps(source_chunk_ids or [], ensure_ascii=False)
        node.properties = json.dumps(properties or {}, ensure_ascii=False)
        db.session.flush()
        return node, is_new

    def _upsert_edge(self, graph_id, course_id, source_id, target_id, edge_type, weight, confidence, evidence_chunk_ids):
        if source_id == target_id:
            return 0
        edge_type = edge_type if edge_type in EDGE_TYPES else "related"
        cache_key = (course_id, source_id, target_id, edge_type)
        edge = self._edge_cache.get(cache_key)
        if edge:
            edge.weight = max(edge.weight or 0, weight)
            edge.confidence = max(edge.confidence or 0, confidence)
            return 0
        new_edge = KnowledgeGraphEdge(
            graph_id=graph_id,
            course_id=course_id,
            source_node_id=source_id,
            target_node_id=target_id,
            edge_type=edge_type,
            weight=weight,
            confidence=confidence,
            evidence_chunk_ids=json.dumps(evidence_chunk_ids or [], ensure_ascii=False),
        )
        db.session.add(new_edge)
        self._edge_cache[cache_key] = new_edge
        return 1

    def _chunk_ids_for_text(self, chunks, text):
        text = (text or "").lower()
        if not text:
            return []
        matches = [c.id for c in chunks if text[:30] in (c.content or "").lower()]
        return matches[:3] or ([chunks[0].id] if chunks else [])

    def _infer_cross_chapter_relations(self, graph_id, course_id, kp_nodes, source_chunks):
        """Infer cross-chapter knowledge relations."""
        created = 0
        # 1. 使用 LLM 返回的跨章节关系
        cross_relations = getattr(self, '_cross_chapter_relations', [])
        if cross_relations:
            # 构建标签 →节点映射
            label_to_node = {}
            for node in kp_nodes:
                label_to_node[node.label.strip()] = node
                # 也支持部分匹配
                label_to_node[node.label.strip().lower()] = node

            for rel in cross_relations:
                source_label = rel.get("source", "").strip()
                target_label = rel.get("target", "").strip()
                relation_type = rel.get("relation", "related")
                strength = rel.get("strength", 0.6)

                source_node = label_to_node.get(source_label) or label_to_node.get(source_label.lower())
                target_node = label_to_node.get(target_label) or label_to_node.get(target_label.lower())

                # 模糊匹配：如果精确匹配失败，尝试包含匹配
                if not source_node:
                    for label, node in label_to_node.items():
                        if source_label in label or label in source_label:
                            source_node = node
                            break
                if not target_node:
                    for label, node in label_to_node.items():
                        if target_label in label or label in target_label:
                            target_node = node
                            break

                if source_node and target_node and source_node.id != target_node.id:
                    edge_type = relation_type if relation_type in EDGE_TYPES else "related"
                    weight = min(1.0, max(0.3, strength))
                    created += self._upsert_edge(
                        graph_id, course_id, source_node.id, target_node.id,
                        edge_type, weight, weight * 0.9, []
                    )

        # 2. 基于关键词共现的关联推理
        if len(kp_nodes) > 1:
            created += self._infer_keyword_relations(graph_id, course_id, kp_nodes)

        return created

    def _infer_keyword_relations(self, graph_id, course_id, kp_nodes):
        """基于关键词重叠推理知识点间的关联（倒排索引优化）"""
        created = 0
        # 提取每个知识点的关键词
        node_keywords = {}
        for node in kp_nodes:
            desc = (node.description or "").lower()
            label = (node.label or "").lower()
            words = set(re.split(r'[，。、；；\s]+', desc + ' ' + label))
            words = {w for w in words if len(w) >= 2}
            node_keywords[node.id] = words

        # 构建倒排索引：keyword → [node_ids]
        inverted_index = {}
        for node_id, keywords in node_keywords.items():
            for kw in keywords:
                inverted_index.setdefault(kw, []).append(node_id)

        # 只比较共享关键词的节点对，限制每节点最多5个关联
        pair_scores = {}
        node_relation_count = {}
        max_relations_per_node = 5

        for kw, node_ids in inverted_index.items():
            if len(node_ids) < 2:
                continue
            for i in range(len(node_ids)):
                for j in range(i + 1, len(node_ids)):
                    a, b = node_ids[i], node_ids[j]
                    if a == b:
                        continue
                    pair_key = (min(a, b), max(a, b))
                    pair_scores.setdefault(pair_key, 0)
                    pair_scores[pair_key] += 1

        # 按共享关键词数量排序，取 top 关联
        sorted_pairs = sorted(pair_scores.items(), key=lambda x: x[1], reverse=True)
        for (a, b), shared_count in sorted_pairs:
            if node_relation_count.get(a, 0) >= max_relations_per_node:
                continue
            if node_relation_count.get(b, 0) >= max_relations_per_node:
                continue
            kw_a, kw_b = node_keywords.get(a, set()), node_keywords.get(b, set())
            if not kw_a or not kw_b:
                continue
            overlap = shared_count
            if overlap >= 2:
                jaccard = overlap / len(kw_a | kw_b)
                if jaccard >= 0.15:
                    weight = min(0.8, 0.3 + jaccard)
                    created += self._upsert_edge(
                        graph_id, course_id, a, b, "related", weight, weight * 0.85, []
                    )
                    node_relation_count[a] = node_relation_count.get(a, 0) + 1
                    node_relation_count[b] = node_relation_count.get(b, 0) + 1
        return created

    def _infer_semantic_relations(self, graph_id, course_id, kp_nodes, source_chunks):
        """Infer applies_to relations between practice/case nodes and core concepts."""
        created = 0
        core_kps = []
        practice_kps = []
        practice_markers = ["实践", "案例", "练习", "实验", "exercise", "case", "practice"]
        for node in kp_nodes:
            cat = (node.category or "").lower()
            label = (node.label or "").lower()
            if any(marker in cat or marker in label for marker in practice_markers):
                practice_kps.append(node)
            else:
                core_kps.append(node)

        for practice_node in practice_kps:
            p_label = (practice_node.label or "").lower()
            p_desc = (practice_node.description or "").lower()
            for core_node in core_kps:
                c_label = (core_node.label or "").lower()
                if c_label and (c_label in p_label or c_label in p_desc):
                    created += self._upsert_edge(
                        graph_id, course_id, practice_node.id, core_node.id,
                        "applies_to", 0.7, 0.75, []
                    )
                    break
        return created

    def _apply_agent_relations(self, graph_id, course_id, kp_nodes, relations):
        node_index = {}
        for node in kp_nodes:
            node_index.setdefault(node.label.strip(), []).append(node)
            node_index.setdefault(node.label.strip().lower(), []).append(node)

        def resolve(label):
            if not label:
                return None
            raw = str(label).strip()
            lower = raw.lower()
            if raw in node_index and node_index[raw]:
                return node_index[raw][0]
            if lower in node_index and node_index[lower]:
                return node_index[lower][0]
            for key, nodes in node_index.items():
                if raw in key or key in raw:
                    return nodes[0]
            return None

        created = 0
        for rel in relations or []:
            source_node = resolve(rel.get("source"))
            target_node = resolve(rel.get("target"))
            if not source_node or not target_node or source_node.id == target_node.id:
                continue
            edge_type = rel.get("relation") if rel.get("relation") in EDGE_TYPES else "related"
            weight = float(rel.get("weight") or 0.6)
            confidence = float(rel.get("confidence") or weight)
            created += self._upsert_edge(
                graph_id,
                course_id,
                source_node.id,
                target_node.id,
                edge_type,
                max(0.3, min(1.0, weight)),
                max(0.3, min(1.0, confidence)),
                [],
            )
        return created

    def _clear_existing_graph(self, course_id, delete_all_chunks=False):
        KnowledgeGraphEdge.query.filter_by(course_id=course_id).delete()
        KnowledgeGraphNode.query.filter_by(course_id=course_id).delete()
        if delete_all_chunks:
            # 彻底清除：删除所有类型的来源片段
            KnowledgeSourceChunk.query.filter_by(course_id=course_id).delete()
        else:
            # 仅清除大纲导入的来源片段
            KnowledgeSourceChunk.query.filter_by(course_id=course_id, source_type="syllabus").delete()
        db.session.flush()

    def _run_with_sqlite_lock_retry(self, func, operation="database write", attempts=5):
        for attempt in range(1, attempts + 1):
            try:
                return func()
            except OperationalError as exc:
                if not self._is_sqlite_locked_error(exc) or attempt >= attempts:
                    raise
                db.session.rollback()
                delay = min(0.2 * (2 ** (attempt - 1)), 2.0)
                logger.warning(
                    "SQLite database locked during %s; retrying in %.1fs (%s/%s)",
                    operation,
                    delay,
                    attempt,
                    attempts,
                )
                time.sleep(delay)

    def _commit_with_sqlite_lock_retry(self, operation="database commit", attempts=5):
        return self._run_with_sqlite_lock_retry(db.session.commit, operation=operation, attempts=attempts)

    def _is_sqlite_locked_error(self, exc):
        message = str(exc).lower()
        return "database is locked" in message or "database table is locked" in message

    def _format_node_for_graph(self, node, source_map=None):
        data = node.to_dict(include_sources=True)
        data["label"] = node.label
        data["type"] = node.node_type
        data["size"] = 18 + min(28, int((node.weight or 1) * 10))
        data["color"] = {
            "course": "#2563eb",
            "chapter": "#059669",
            "knowledge_point": "#d97706",
            "objective": "#7c3aed",
            "skill": "#dc2626",
            "case": "#0891b2",
            "exercise": "#4f46e5",
            "resource": "#475569",
        }.get(node.node_type, "#64748b")
        if source_map is not None:
            data["sources"] = [source_map[cid] for cid in data.get("source_chunk_ids", []) if cid in source_map]
        # 为知识点节点添加内容摘要summary字段，供前端3D节点直接显示
        if node.node_type == 'knowledge_point':
            props = data.get('properties', {})
            if isinstance(props, str):
                try:
                    props = json.loads(props)
                except (json.JSONDecodeError, TypeError):
                    props = {}
            desc = node.description or ''
            summary = props.get('summary', '') if isinstance(props, dict) else ''
            # 优先使用description，其次summary，截取前80字作为摘要
            data['summary'] = (desc or summary)[:80] if (desc or summary) else ''
        return data

    def _build_metrics(self, nodes, edges):
        by_type = defaultdict(int)
        for node in nodes:
            by_type[node.node_type] += 1
        return {
            "node_count": len(nodes),
            "edge_count": len(edges),
            "node_types": dict(by_type),
            "average_weight": round(sum(e.weight or 0 for e in edges) / max(len(edges), 1), 2),
        }


syllabus_graph_service = SyllabusGraphService()
