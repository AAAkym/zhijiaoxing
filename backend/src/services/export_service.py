import io
import logging
import os
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

FONT_DIR = Path(os.environ.get("SYSTEMROOT", r"C:\Windows")) / "Fonts"
SIMHEI_PATH = FONT_DIR / "simhei.ttf"
MSYH_PATH = FONT_DIR / "msyh.ttc"

COLOR_RED = (220, 50, 50)
COLOR_GREEN = (22, 101, 52)
COLOR_PURPLE = (128, 0, 128)
COLOR_GRAY = (128, 128, 128)
COLOR_BLACK = (0, 0, 0)
COLOR_DARK = (51, 51, 51)
COLOR_LIGHT_GRAY = (200, 200, 200)

STATUS_MAP_ZH = {"unmastered": "未掌握", "reviewing": "复习中", "mastered": "已掌握"}

ERROR_TYPE_MAP = {
    "concept_understanding": "概念理解偏差",
    "calculation_error": "计算失误",
    "question_misread": "审题不清",
    "careless": "粗心失误",
    "other": "其他",
}

_cached_font_name = None


def _get_cached_font_name():
    global _cached_font_name
    if _cached_font_name is not None:
        return _cached_font_name
    if SIMHEI_PATH.exists():
        _cached_font_name = "SimHei"
    elif MSYH_PATH.exists():
        _cached_font_name = "MSYH"
    else:
        _cached_font_name = "Helvetica"
    return _cached_font_name


class MistakePDF(FPDF):
    def __init__(self):
        super().__init__(orientation="P", unit="mm", format="A4")
        self.set_auto_page_break(auto=True, margin=15)
        self.set_margins(15, 20, 15)
        self._usable_width = self.w - self.l_margin - self.r_margin
        self._fonts_registered = False
        self._font_name = _get_cached_font_name()

    def _register_fonts(self):
        if self._fonts_registered:
            return
        try:
            if self._font_name == "SimHei" and SIMHEI_PATH.exists():
                self.add_font("SimHei", fname=str(SIMHEI_PATH))
                self.add_font("SimHei", style="B", fname=str(SIMHEI_PATH))
            elif self._font_name == "MSYH" and MSYH_PATH.exists():
                self.add_font("MSYH", fname=str(MSYH_PATH))
                self.add_font("MSYH", style="B", fname=str(MSYH_PATH))
        except Exception as e:
            logger.warning(f"Font registration warning: {e}")
        self._fonts_registered = True

    def header(self):
        if self.page_no() == 1:
            return
        self.set_font(self._font_name, "B", 9)
        self.set_text_color(*COLOR_GRAY)
        self.cell(0, 8, "错题本导出", align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def footer(self):
        self.set_y(-12)
        self.set_font(self._font_name, "", 8)
        self.set_text_color(*COLOR_GRAY)
        self.cell(0, 8, f"第 {self.page_no()} 页", align="C")

    def _write_title_page(self, total_count, export_mode="full"):
        self._register_fonts()
        fn = self._font_name

        self.set_font(fn, "B", 22)
        self.set_text_color(*COLOR_DARK)
        self.ln(30)
        self.cell(0, 15, "错题本导出", align="C", new_x="LMARGIN", new_y="NEXT")

        self.set_font(fn, "", 11)
        self.set_text_color(*COLOR_GRAY)
        self.cell(0, 8, f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}", align="C", new_x="LMARGIN", new_y="NEXT")
        self.cell(0, 8, f"共 {total_count} 道错题", align="C", new_x="LMARGIN", new_y="NEXT")

        mode_labels = {"full": "完整导出", "questions_only": "仅题干模式（重新练习用）"}
        self.cell(0, 8, f"导出模式：{mode_labels.get(export_mode, export_mode)}", align="C", new_x="LMARGIN", new_y="NEXT")

        self.ln(20)
        self.set_draw_color(*COLOR_LIGHT_GRAY)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(10)

    def _write_label_value(self, label, value, font_size=10, color=None, bold_label=True, indent=0):
        if not value and value != 0:
            return
        value = str(value).strip()
        if not value:
            return

        fn = self._font_name
        self.set_x(self.l_margin + indent)

        if label:
            self.set_font(fn, "B" if bold_label else "", font_size)
            self.set_text_color(*COLOR_DARK)
            label_w = self.get_string_width(label) + 2
            min_label_w = self._usable_width * 0.25
            if label_w < min_label_w:
                label_w = min_label_w
            self.cell(label_w - indent, 7, label, new_x="END", new_y="TOP")

        self.set_font(fn, "", font_size)
        self.set_text_color(*(color if color else COLOR_BLACK))
        remaining_w = self.w - self.r_margin - self.get_x()
        if remaining_w < 20:
            self.ln(7)
            self.set_x(self.l_margin + indent)
            remaining_w = self._usable_width - indent
        self.multi_cell(remaining_w, 7, value, new_x="LMARGIN", new_y="NEXT")

    def _write_multiline(self, text, font_size=10, color=None, style="", indent=0):
        if not text:
            return
        text = str(text).strip()
        if not text:
            return

        fn = self._font_name
        self.set_font(fn, style, font_size)
        self.set_text_color(*(color if color else COLOR_BLACK))
        self.set_x(self.l_margin + indent)
        w = self._usable_width - indent
        self.multi_cell(w, 7, text, new_x="LMARGIN", new_y="NEXT")

    def _needs_new_page(self, needed_mm=40):
        return self.get_y() + needed_mm > self.h - self.b_margin

    def add_mistake_detailed(self, index, m, export_mode="full"):
        fn = self._font_name

        if self._needs_new_page(50):
            self.add_page()

        self.set_font(fn, "B", 13)
        self.set_text_color(*COLOR_DARK)
        self.cell(0, 10, f"第 {index} 题", new_x="LMARGIN", new_y="NEXT")

        course_title = m.get("course_title", "")
        if course_title:
            self._write_label_value("所属课程：", course_title, font_size=10, color=COLOR_GRAY)

        assessment_title = m.get("assessment_title", "")
        if assessment_title:
            self._write_label_value("所属考核：", assessment_title, font_size=10, color=COLOR_GRAY)

        self._write_label_value("题目内容：", m.get("question_content", ""), font_size=11)

        options = m.get("options", [])
        if options and isinstance(options, list):
            for idx, opt in enumerate(options):
                label = chr(65 + idx)
                opt_text = opt if isinstance(opt, str) else str(opt)
                self._write_multiline(f"{label}. {opt_text}", font_size=10, indent=8)

        if export_mode == "questions_only":
            self.ln(3)
            y = self.get_y()
            if y < self.h - 20:
                self.set_draw_color(*COLOR_LIGHT_GRAY)
                self.line(self.l_margin, y, self.w - self.r_margin, y)
            self.ln(5)
            return

        user_ans = m.get("user_answer_display") or m.get("user_answer", "")
        self._write_label_value("你的答案：", user_ans, font_size=11, color=COLOR_RED)

        correct_ans = m.get("correct_answer_display") or m.get("correct_answer", "")
        self._write_label_value("正确答案：", correct_ans, font_size=11, color=COLOR_GREEN)

        mastery = m.get("mastery_status", "unmastered")
        status_text = STATUS_MAP_ZH.get(mastery, mastery)
        self._write_label_value("掌握状态：", status_text, font_size=10)

        mistake_count = m.get("mistake_count", 1)
        if mistake_count > 1:
            self._write_label_value("错误次数：", str(mistake_count), font_size=10)

        tags = m.get("knowledge_tags", [])
        if tags:
            tags_str = "、".join(tags) if isinstance(tags, list) else str(tags)
            self._write_label_value("知识点标签：", tags_str, font_size=10)

        error_type = m.get("error_type", "")
        if error_type:
            error_text = ERROR_TYPE_MAP.get(error_type, error_type)
            self._write_label_value("错误类型：", error_text, font_size=10)

        ai_analysis = m.get("ai_analysis", "")
        if ai_analysis:
            self._write_label_value("AI 分析：", ai_analysis, font_size=10, color=COLOR_PURPLE)

        error_detail = m.get("error_reason_detail", "")
        if error_detail:
            self._write_label_value("错因详情：", error_detail, font_size=10)

        explanation = ""
        orig = m.get("original_question")
        if orig and isinstance(orig, dict):
            explanation = orig.get("explanation", "")
        if explanation:
            self._write_label_value("题目解析：", explanation, font_size=10, color=COLOR_GREEN)

        self.ln(3)
        y = self.get_y()
        if y < self.h - 20:
            self.set_draw_color(*COLOR_LIGHT_GRAY)
            self.line(self.l_margin, y, self.w - self.r_margin, y)
        self.ln(5)

    def add_mistake_compact(self, index, m, export_mode="full"):
        fn = self._font_name
        self.set_font(fn, "", 9)
        self.set_text_color(*COLOR_BLACK)

        question = m.get("question_content", "")[:60]
        mastery = m.get("mastery_status", "unmastered")
        status = STATUS_MAP_ZH.get(mastery, mastery)

        if export_mode == "questions_only":
            line = f"{index}. {question}"
        else:
            user_ans = m.get("user_answer_display") or m.get("user_answer", "")
            correct_ans = m.get("correct_answer_display") or m.get("correct_answer", "")
            line = f"{index}. {question} | 你的答案: {user_ans} | 正确答案: {correct_ans} | {status}"

        self.set_x(self.l_margin)
        self.multi_cell(self._usable_width, 6, line, new_x="LMARGIN", new_y="NEXT")


def generate_pdf(
    mistakes_data: List[Dict],
    template: str = "detailed",
    export_mode: str = "full",
) -> bytes:
    t0 = time.perf_counter()
    pdf = MistakePDF()
    pdf._register_fonts()
    pdf.add_page()
    pdf._write_title_page(len(mistakes_data), export_mode=export_mode)

    if template == "compact":
        for i, m in enumerate(mistakes_data, 1):
            pdf.add_mistake_compact(i, m, export_mode=export_mode)
    else:
        for i, m in enumerate(mistakes_data, 1):
            pdf.add_mistake_detailed(i, m, export_mode=export_mode)

    result = pdf.output()
    elapsed = time.perf_counter() - t0
    logger.info(f"PDF generation: {len(mistakes_data)} mistakes, template={template}, mode={export_mode}, {elapsed:.3f}s, {len(result)//1024}KB")
    return result


def generate_docx(
    mistakes_data: List[Dict],
    template: str = "detailed",
    export_mode: str = "full",
) -> bytes:
    t0 = time.perf_counter()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    style.font.size = Pt(11)

    title = doc.add_heading("错题本导出", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(*COLOR_DARK)

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mode_labels = {"full": "完整导出", "questions_only": "仅题干模式（重新练习用）"}
    run = info_para.add_run(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}  |  共 {len(mistakes_data)} 道错题  |  {mode_labels.get(export_mode, export_mode)}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(*COLOR_GRAY)

    doc.add_paragraph()

    if template == "compact":
        cols = 3 if export_mode == "questions_only" else 5
        table = doc.add_table(rows=1, cols=cols)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "序号"
        hdr[1].text = "题目"
        if export_mode == "full":
            hdr[2].text = "你的答案"
            hdr[3].text = "正确答案"
            hdr[4].text = "状态"
        else:
            hdr[2].text = "状态"
        for i, m in enumerate(mistakes_data, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = m.get("question_content", "")[:60]
            if export_mode == "full":
                row[2].text = m.get("user_answer_display") or m.get("user_answer", "")
                row[3].text = m.get("correct_answer_display") or m.get("correct_answer", "")
                mastery = m.get("mastery_status", "unmastered")
                row[4].text = STATUS_MAP_ZH.get(mastery, mastery)
            else:
                mastery = m.get("mastery_status", "unmastered")
                row[2].text = STATUS_MAP_ZH.get(mastery, mastery)
        buffer = io.BytesIO()
        doc.save(buffer)
        elapsed = time.perf_counter() - t0
        logger.info(f"DOCX compact generation: {len(mistakes_data)} mistakes, mode={export_mode}, {elapsed:.3f}s")
        return buffer.getvalue()

    for i, m in enumerate(mistakes_data, 1):
        heading = doc.add_heading(f"第 {i} 题", level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*COLOR_DARK)

        course_title = m.get("course_title", "")
        if course_title:
            p = doc.add_paragraph()
            run_label = p.add_run("所属课程：")
            run_label.bold = True
            run_label.font.size = Pt(10)
            run_label.font.color.rgb = RGBColor(*COLOR_GRAY)
            p.add_run(course_title).font.size = Pt(10)

        assessment_title = m.get("assessment_title", "")
        if assessment_title:
            p = doc.add_paragraph()
            run_label = p.add_run("所属考核：")
            run_label.bold = True
            run_label.font.size = Pt(10)
            run_label.font.color.rgb = RGBColor(*COLOR_GRAY)
            p.add_run(assessment_title).font.size = Pt(10)

        p = doc.add_paragraph()
        run_label = p.add_run("题目内容：")
        run_label.bold = True
        p.add_run(m.get("question_content", ""))

        options = m.get("options", [])
        if options and isinstance(options, list):
            for idx, opt in enumerate(options):
                label = chr(65 + idx)
                opt_text = opt if isinstance(opt, str) else str(opt)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.add_run(f"{label}. {opt_text}").font.size = Pt(10)

        if export_mode == "questions_only":
            doc.add_paragraph()
            continue

        p = doc.add_paragraph()
        run_label = p.add_run("你的答案：")
        run_label.bold = True
        user_ans = m.get("user_answer_display") or m.get("user_answer", "")
        run_ans = p.add_run(user_ans)
        run_ans.font.color.rgb = RGBColor(*COLOR_RED)

        p = doc.add_paragraph()
        run_label = p.add_run("正确答案：")
        run_label.bold = True
        correct_ans = m.get("correct_answer_display") or m.get("correct_answer", "")
        run_ans = p.add_run(correct_ans)
        run_ans.font.color.rgb = RGBColor(*COLOR_GREEN)

        mastery = m.get("mastery_status", "unmastered")
        p = doc.add_paragraph()
        run_label = p.add_run("掌握状态：")
        run_label.bold = True
        p.add_run(STATUS_MAP_ZH.get(mastery, mastery))

        mistake_count = m.get("mistake_count", 1)
        if mistake_count > 1:
            p = doc.add_paragraph()
            run_label = p.add_run("错误次数：")
            run_label.bold = True
            p.add_run(str(mistake_count))

        tags = m.get("knowledge_tags", [])
        if tags:
            tags_str = "、".join(tags) if isinstance(tags, list) else str(tags)
            p = doc.add_paragraph()
            run_label = p.add_run("知识点标签：")
            run_label.bold = True
            p.add_run(tags_str)

        error_type = m.get("error_type", "")
        if error_type:
            p = doc.add_paragraph()
            run_label = p.add_run("错误类型：")
            run_label.bold = True
            p.add_run(ERROR_TYPE_MAP.get(error_type, error_type))

        ai_analysis = m.get("ai_analysis", "")
        if ai_analysis:
            p = doc.add_paragraph()
            run_label = p.add_run("AI 分析：")
            run_label.bold = True
            run_label.font.color.rgb = RGBColor(*COLOR_PURPLE)
            p.add_run(ai_analysis)

        error_detail = m.get("error_reason_detail", "")
        if error_detail:
            p = doc.add_paragraph()
            run_label = p.add_run("错因详情：")
            run_label.bold = True
            p.add_run(error_detail)

        explanation = ""
        orig = m.get("original_question")
        if orig and isinstance(orig, dict):
            explanation = orig.get("explanation", "")
        if explanation:
            p = doc.add_paragraph()
            run_label = p.add_run("题目解析：")
            run_label.bold = True
            run_label.font.color.rgb = RGBColor(*COLOR_GREEN)
            p.add_run(explanation)

        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    elapsed = time.perf_counter() - t0
    logger.info(f"DOCX detailed generation: {len(mistakes_data)} mistakes, mode={export_mode}, {elapsed:.3f}s")
    return buffer.getvalue()
